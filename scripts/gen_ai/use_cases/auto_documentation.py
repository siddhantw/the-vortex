import streamlit as st
import pandas as pd
import numpy as np
import os
import sys
import json
import re
import time
from datetime import datetime, timedelta
from io import StringIO, BytesIO
from pathlib import Path
import statistics
import base64  # Added for embedding images in HTML

# Ensure the parent directory is in the path to import shared modules
parent_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
if parent_dir not in sys.path:
    sys.path.append(parent_dir)

# Import notifications module for sending notifications
try:
    import notifications
    NOTIFICATIONS_AVAILABLE = True
except ImportError:
    NOTIFICATIONS_AVAILABLE = False

# Define supported output formats
OUTPUT_FORMATS = {
    "HTML": {"extension": "html", "mime": "text/html"},
    "Markdown": {"extension": "md", "mime": "text/markdown"},
    "PDF": {"extension": "pdf", "mime": "application/pdf"},
    "Word": {"extension": "docx", "mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"},
}

# Check for optional dependencies
try:
    import matplotlib.pyplot as plt
    import seaborn as sns
    VISUALIZATION_AVAILABLE = True
except ImportError:
    VISUALIZATION_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    PDF_GENERATION_AVAILABLE = True
except ImportError:
    PDF_GENERATION_AVAILABLE = False

try:
    from docx import Document
    DOCX_GENERATION_AVAILABLE = True
except ImportError:
    DOCX_GENERATION_AVAILABLE = False

# Check for AI analysis dependencies
try:
    import nltk
    from nltk.tokenize import word_tokenize
    from nltk.corpus import stopwords
    # Download required NLTK resources
    try:
        nltk.data.find('tokenizers/punkt')
    except LookupError:
        nltk.download('punkt', quiet=True)
    try:
        nltk.data.find('corpora/stopwords')
    except LookupError:
        nltk.download('stopwords', quiet=True)
    AI_ANALYSIS_AVAILABLE = True
except ImportError:
    AI_ANALYSIS_AVAILABLE = False

class TestMetricsCalculator:
    """Helper class to calculate advanced test metrics from parsed test data."""

    @staticmethod
    def calculate_coverage_metrics(data):
        """Calculate test coverage metrics based on available data."""
        metrics = {
            "functional_coverage": 0,
            "line_coverage": 0,
            "branch_coverage": 0,
            "api_coverage": 0,
            "ui_coverage": 0,
            "has_coverage_data": False
        }

        # Look for coverage info in test data
        if "coverage" in data:
            metrics["has_coverage_data"] = True
            coverage = data["coverage"]
            metrics.update(coverage)

        # Estimate functional coverage based on test count and pass rate if no actual data
        if metrics["functional_coverage"] == 0:
            test_count = data['summary']['total']
            if test_count > 100:
                metrics["functional_coverage"] = min(90, test_count / 10)  # Estimate based on test count
            elif test_count > 50:
                metrics["functional_coverage"] = min(80, test_count / 8)
            elif test_count > 20:
                metrics["functional_coverage"] = min(70, test_count / 5)
            else:
                metrics["functional_coverage"] = min(60, test_count * 2)

        return metrics

    @staticmethod
    def calculate_flakiness_metrics(data):
        """Calculate test flakiness metrics."""
        metrics = {
            "flaky_tests": 0,
            "flaky_rate": 0,
            "flaky_test_names": [],
            "has_flakiness_data": False,
            "flakiness_score": 100  # Higher is better (less flaky)
        }

        # Look for flakiness info in test data
        if "flakiness" in data:
            metrics["has_flakiness_data"] = True
            metrics.update(data["flakiness"])
        else:
            # Attempt to detect flaky tests based on available data
            unstable_indicators = ['intermittent', 'flaky', 'unstable', 'randomly', 'sometimes']

            flaky_tests = []
            for test in data['tests']:
                # Check test name and messages for flakiness indicators
                test_text = f"{test.get('name', '')} {test.get('message', '')}"
                if any(indicator in test_text.lower() for indicator in unstable_indicators):
                    flaky_tests.append(test.get('name', 'Unknown Test'))

            if flaky_tests:
                metrics["flaky_tests"] = len(flaky_tests)
                metrics["flaky_rate"] = len(flaky_tests) / data['summary']['total'] * 100
                metrics["flaky_test_names"] = flaky_tests
                metrics["flakiness_score"] = 100 - metrics["flaky_rate"]
                metrics["has_flakiness_data"] = True

        return metrics

    @staticmethod
    def calculate_quality_metrics(data):
        """Calculate overall test quality metrics."""
        metrics = {
            "test_density": 0,  # Tests per feature/unit
            "complexity": 0,    # Test complexity score
            "stability": 0,     # Test stability score
            "avg_maintenance_cost": 0  # Estimated maintenance cost
        }

        # Calculate test density if we have feature/component info
        component_count = {}
        for test in data['tests']:
            component = None
            # Try to find component/class/feature information
            if 'classname' in test:
                parts = test.get('classname', '').split('.')
                component = parts[0] if parts else None
            elif 'feature' in test:
                component = test.get('feature')

            if component:
                component_count[component] = component_count.get(component, 0) + 1

        if component_count:
            # Calculate average tests per component/feature
            metrics["test_density"] = sum(component_count.values()) / len(component_count)

        # Calculate test complexity based on test durations
        if data['tests'] and 'time' in data['tests'][0]:
            test_times = [t.get('time', 0) for t in data['tests']]
            if test_times:
                avg_time = sum(test_times) / len(test_times)
                time_variance = sum((t - avg_time) ** 2 for t in test_times) / len(test_times)

                # Complexity formula: combines average time and variance
                metrics["complexity"] = min(100, (avg_time * 10) + (time_variance * 5))

        # Calculate stability score based on pass rate and failed test variance
        pass_rate = data['summary']['passed'] / data['summary']['total'] * 100 if data['summary']['total'] > 0 else 0
        metrics["stability"] = pass_rate

        # Estimate maintenance cost
        if data['summary']['failed'] > 0:
            # Higher maintenance cost with more failures and skips
            failure_rate = data['summary']['failed'] / data['summary']['total'] if data['summary']['total'] > 0 else 0
            skip_rate = data['summary']['skipped'] / data['summary']['total'] if data['summary']['total'] > 0 else 0
            metrics["avg_maintenance_cost"] = (failure_rate * 80) + (skip_rate * 40) + 10
        else:
            # Base maintenance cost
            metrics["avg_maintenance_cost"] = 10 + (data['summary']['skipped'] * 2)

        return metrics

    @staticmethod
    def calculate_all_metrics(data):
        """Calculate all available metrics from test data."""
        all_metrics = {
            "coverage": TestMetricsCalculator.calculate_coverage_metrics(data),
            "flakiness": TestMetricsCalculator.calculate_flakiness_metrics(data),
            "quality": TestMetricsCalculator.calculate_quality_metrics(data)
        }

        # Add high-level composite scores
        all_metrics["overall_health"] = min(100, (
            all_metrics["coverage"]["functional_coverage"] * 0.4 +
            all_metrics["flakiness"]["flakiness_score"] * 0.4 +
            all_metrics["quality"]["stability"] * 0.2
        ))

        all_metrics["maintainability"] = max(0, 100 - all_metrics["quality"]["avg_maintenance_cost"])

        return all_metrics

def parse_junit_xml(file_content):
    """Parse JUnit XML test results into structured data."""
    import xml.etree.ElementTree as ET

    try:
        root = ET.fromstring(file_content)

        # Initialize data structure
        data = {
            "tests": [],
            "summary": {
                "total": 0,
                "passed": 0,
                "failed": 0,
                "skipped": 0,
                "execution_time": 0
            }
        }

        # First try to get total time from the root testsuites element if available
        if root.tag == 'testsuites':
            root_time_str = root.attrib.get('time', None)
            if root_time_str:
                try:
                    data["summary"]["execution_time"] = float(root_time_str)
                except (ValueError, TypeError):
                    pass

        # Find all testcase elements
        for testsuite in root.findall(".//testsuite"):
            suite_name = testsuite.attrib.get('name', 'Unknown Suite')

            # Parse time carefully to handle different formats and potential missing values
            suite_time_str = testsuite.attrib.get('time', '0')
            try:
                suite_time = float(suite_time_str)
            except (ValueError, TypeError):
                suite_time = 0

            # Accumulate total execution time from suites if we don't have it from the root
            if data["summary"]["execution_time"] == 0:
                data["summary"]["execution_time"] += suite_time

            # Check for timestamp attributes for more accurate timing
            if 'timestamp' in testsuite.attrib:
                try:
                    timestamp = datetime.fromisoformat(testsuite.attrib['timestamp'].replace('Z', '+00:00'))
                    if hasattr(data, 'start_time'):
                        data.start_time = min(data.start_time, timestamp)
                    else:
                        data.start_time = timestamp
                except (ValueError, TypeError):
                    pass

            for testcase in testsuite.findall(".//testcase"):
                # Parse time carefully to avoid issues with bad time values
                test_time_str = testcase.attrib.get('time', '0')
                try:
                    test_time = float(test_time_str)
                except (ValueError, TypeError):
                    test_time = 0

                test_data = {
                    "name": testcase.attrib.get('name', 'Unknown'),
                    "classname": testcase.attrib.get('classname', ''),
                    "time": test_time,
                    "status": "passed"
                }

                # Check for failures or errors
                failure = testcase.find('failure')
                error = testcase.find('error')
                skipped = testcase.find('skipped')

                if failure is not None:
                    test_data["status"] = "failed"
                    test_data["message"] = failure.attrib.get('message', '')
                    test_data["type"] = failure.attrib.get('type', '')
                    test_data["stacktrace"] = failure.text
                    data["summary"]["failed"] += 1
                elif error is not None:
                    test_data["status"] = "error"
                    test_data["message"] = error.attrib.get('message', '')
                    test_data["type"] = error.attrib.get('type', '')
                    test_data["stacktrace"] = error.text
                    data["summary"]["failed"] += 1
                elif skipped is not None:
                    test_data["status"] = "skipped"
                    test_data["message"] = skipped.attrib.get('message', '')
                    data["summary"]["skipped"] += 1
                else:
                    data["summary"]["passed"] += 1

                data["tests"].append(test_data)
                data["summary"]["total"] += 1

        # If no execution_time was captured from suite times, calculate it from test times
        if data["summary"]["execution_time"] == 0 and data["tests"]:
            total_test_time = sum(test.get("time", 0) for test in data["tests"])
            if total_test_time > 0:
                data["summary"]["execution_time"] = total_test_time

        # Ensure we have a non-zero execution time
        if data["summary"]["execution_time"] == 0 and data["summary"]["total"] > 0:
            # If all else fails, estimate based on test count
            data["summary"]["execution_time"] = data["summary"]["total"] * 0.5  # Assume 0.5s per test as fallback

        return data

    except Exception as e:
        st.error(f"Error parsing JUnit XML: {e}")
        return None

def parse_robot_xml(file_content):
    """Parse Robot Framework XML output into structured data.

    The Robot Framework output.xml contains all test execution details including
    test counts and execution times. This function parses that information into
    a structured format for report generation.
    """
    import xml.etree.ElementTree as ET
    import re

    try:
        root = ET.fromstring(file_content)

        # Initialize data structure
        data = {
            "tests": [],
            "summary": {
                "total": 0,
                "passed": 0,
                "failed": 0,
                "skipped": 0,
                "execution_time": 0
            },
            "suites": []
        }

        # Extract execution time directly from statistics which is most accurate
        statistics = root.find('.//statistics')
        if statistics is not None:
            total_stats = statistics.find('total')
            if total_stats is not None:
                all_stats = total_stats.findall('stat')
                for stat in all_stats:
                    if 'All Tests' in stat.text:
                        if 'elapsed' in stat.attrib:
                            try:
                                data["summary"]["execution_time"] = float(stat.attrib.get('elapsed', 0)) / 1000
                                break
                            except (ValueError, TypeError):
                                pass

        # If statistics didn't have time, try other methods
        if data["summary"]["execution_time"] == 0:
            # Extract from root if available (most reliable after statistics)
            if 'elapsedtime' in root.attrib:
                try:
                    data["summary"]["execution_time"] = float(root.attrib['elapsedtime']) / 1000
                except (ValueError, TypeError):
                    pass

            # Try calculating from timestamps if no elapsedtime
            elif 'endtime' in root.attrib and 'starttime' in root.attrib:
                try:
                    # Robot format: 20250604 12:34:56.789
                    start_str = root.attrib['starttime']
                    end_str = root.attrib['endtime']

                    # Support both formats with and without milliseconds
                    start_time = None
                    end_time = None

                    # Try with milliseconds
                    try:
                        start_time = datetime.strptime(start_str, '%Y%m%d %H:%M:%S.%f')
                        end_time = datetime.strptime(end_str, '%Y%m%d %H:%M:%S.%f')
                    except ValueError:
                        # Try without milliseconds
                        try:
                            start_time = datetime.strptime(start_str, '%Y%m%d %H:%M:%S')
                            end_time = datetime.strptime(end_str, '%Y%m%d %H:%M:%S')
                        except ValueError:
                            pass

                    if start_time and end_time:
                        data["summary"]["execution_time"] = (end_time - start_time).total_seconds()
                except Exception:
                    pass

        # Process all test cases in the XML, traversing the suite structure
        suite_times_total = 0  # For verification
        test_times_total = 0    # For verification

        # Find all test cases regardless of suite structure
        all_tests = root.findall('.//test')

        # Update the total tests count directly based on the XML
        data["summary"]["total"] = len(all_tests)

        # Process all suites
        for suite in root.findall('.//suite'):
            suite_name = suite.attrib.get('name', 'Unknown Suite')
            suite_data = {
                "name": suite_name,
                "tests": []
            }

            # Extract suite elapsed time if available
            if 'elapsedtime' in suite.attrib:
                try:
                    suite_elapsed_ms = float(suite.attrib['elapsedtime'])
                    suite_times_total += suite_elapsed_ms / 1000
                except (ValueError, TypeError):
                    pass

            # Process all test cases in this suite
            for test in suite.findall('./test'):
                test_name = test.attrib.get('name', 'Unknown Test')

                # Get accurate test time from status element (most reliable)
                test_time = 0
                status_elem = test.find('./status')
                if status_elem is not None and 'elapsedtime' in status_elem.attrib:
                    try:
                        test_time = float(status_elem.attrib['elapsedtime']) / 1000
                    except (ValueError, TypeError):
                        # Fall back to test elapsedtime if status elapsedtime fails
                        if 'elapsedtime' in test.attrib:
                            try:
                                test_time = float(test.attrib['elapsedtime']) / 1000
                            except (ValueError, TypeError):
                                test_time = 0
                else:
                    # No status element, try test element directly
                    if 'elapsedtime' in test.attrib:
                        try:
                            test_time = float(test.attrib['elapsedtime']) / 1000
                        except (ValueError, TypeError):
                            test_time = 0

                # Accumulate total test time for verification
                test_times_total += test_time

                # Get status information
                status = "unknown"
                message = ""
                if status_elem is not None:
                    status = status_elem.attrib.get('status', '').lower()
                    message = status_elem.text or ""

                test_data = {
                    "name": test_name,
                    "id": test.attrib.get('id', ''),
                    "time": test_time,
                    "status": status,
                    "message": message
                }

                # Update counters based on status
                if status == 'pass':
                    data["summary"]["passed"] += 1
                elif status == 'fail':
                    data["summary"]["failed"] += 1
                else:  # not run, skipped, etc.
                    data["summary"]["skipped"] += 1

                suite_data["tests"].append(test_data)
                data["tests"].append(test_data)

            data["suites"].append(suite_data)

        # Verify that we have accurate counts
        total_count_check = data["summary"]["passed"] + data["summary"]["failed"] + data["summary"]["skipped"]
        if total_count_check != data["summary"]["total"] and total_count_check > 0:
            # If calculated count doesn't match total but is non-zero, use calculated count
            data["summary"]["total"] = total_count_check

        # Final fallback for execution time if still zero
        if data["summary"]["execution_time"] == 0:
            # If no execution time found in metadata, use suite times or test times
            if suite_times_total > 0:
                data["summary"]["execution_time"] = suite_times_total
            elif test_times_total > 0:
                data["summary"]["execution_time"] = test_times_total

        # Ensure execution time is never zero to avoid display issues
        if data["summary"]["execution_time"] == 0 and data["summary"]["total"] > 0:
            # Last resort: estimate based on test count
            data["summary"]["execution_time"] = data["summary"]["total"] * 0.1  # Conservative fallback estimate

        return data

    except Exception as e:
        st.error(f"Error parsing Robot Framework XML: {e}")
        return None

def parse_cucumber_json(file_content):
    """Parse Cucumber JSON report into structured data."""
    try:
        report = json.loads(file_content)

        # Initialize data structure
        data = {
            "tests": [],
            "summary": {
                "total": 0,
                "passed": 0,
                "failed": 0,
                "skipped": 0,
                "execution_time": 0
            },
            "features": []
        }

        start_time = None
        end_time = None

        # Track start and end times from scenario timestamps if available
        for feature in report:
            feature_data = {
                "name": feature.get("name", "Unknown Feature"),
                "description": feature.get("description", ""),
                "scenarios": []
            }

            for element in feature.get("elements", []):
                if element.get("type") == "scenario":
                    scenario = {
                        "name": element.get("name", "Unknown Scenario"),
                        "steps": [],
                        "status": "passed",  # Default status
                        "time": 0  # Track scenario time
                    }

                    # Try to extract start time if available
                    if "start_timestamp" in element:
                        try:
                            scenario_start = datetime.fromisoformat(element["start_timestamp"].replace('Z', '+00:00'))
                            if start_time is None or scenario_start < start_time:
                                start_time = scenario_start
                        except (ValueError, TypeError):
                            pass

                    # Process steps
                    for step in element.get("steps", []):
                        step_data = {
                            "name": step.get("name", "Unknown Step"),
                            "keyword": step.get("keyword", ""),
                            "status": step.get("result", {}).get("status", "unknown")
                        }

                        # Get execution time with careful error handling
                        try:
                            # Cucumber can use either nanoseconds or milliseconds
                            duration = step.get("result", {}).get("duration", 0)

                            # Detect if duration is in nanoseconds (typically very large numbers)
                            if duration > 1e9:  # If > 1 billion, assume nanoseconds
                                step_time = duration / 1e9  # Convert nanoseconds to seconds
                            else:
                                # Otherwise assume milliseconds
                                step_time = duration / 1000.0  # Convert milliseconds to seconds
                        except (ValueError, TypeError):
                            step_time = 0

                        step_data["time"] = step_time
                        scenario["time"] += step_time

                        # Track end time implicitly by adding duration to start time
                        if start_time is not None:
                            potential_end = start_time + timedelta(seconds=step_time)
                            if end_time is None or potential_end > end_time:
                                end_time = potential_end

                        # Check status
                        if step_data["status"] == "failed":
                            scenario["status"] = "failed"
                            step_data["error_message"] = step.get("result", {}).get("error_message", "")
                        elif step_data["status"] == "skipped" and scenario["status"] != "failed":
                            scenario["status"] = "skipped"

                        scenario["steps"].append(step_data)

                    feature_data["scenarios"].append(scenario)

                    # Add test to overall list and update summary
                    test_data = {
                        "name": scenario["name"],
                        "feature": feature_data["name"],
                        "status": scenario["status"],
                        "steps": len(scenario["steps"]),
                        "time": scenario["time"]  # Add scenario time to test data
                    }
                    data["tests"].append(test_data)
                    data["summary"]["total"] += 1

                    # Add to total execution time only if we aren't using start/end timestamps
                    if start_time is None or end_time is None:
                        data["summary"]["execution_time"] += scenario["time"]

                    if scenario["status"] == "passed":
                        data["summary"]["passed"] += 1
                    elif scenario["status"] == "failed":
                        data["summary"]["failed"] += 1
                    else:
                        data["summary"]["skipped"] += 1

            data["features"].append(feature_data)

        # If we captured start and end times, calculate total execution time
        if start_time and end_time:
            data["summary"]["execution_time"] = (end_time - start_time).total_seconds()

        # If execution time is still 0 but we have tests, ensure we have a value
        if data["summary"]["execution_time"] == 0 and data["tests"]:
            # Double check by summing test times
            total_test_time = sum(test.get("time", 0) for test in data["tests"])
            if total_test_time > 0:
                data["summary"]["execution_time"] = total_test_time
            elif data["summary"]["total"] > 0:
                # Last resort - estimate based on number of tests
                data["summary"]["execution_time"] = data["summary"]["total"] * 0.5  # Assume 0.5s per test as fallback

        return data

    except Exception as e:
        st.error(f"Error parsing Cucumber JSON: {e}")
        return None

def analyze_test_times(data):
    """
    Analyze test execution times and provide insights.
    Returns key metrics and observations about test performance.
    """
    if not data['tests'] or 'time' not in data['tests'][0]:
        return {
            "average_time": 0,
            "median_time": 0,
            "slowest_tests": [],
            "fastest_tests": [],
            "total_time": data['summary']['execution_time'],
            "insights": ["No detailed timing information available"]
        }

    # Extract test times and calculate statistics
    test_times = [(t.get('name', 'Unknown'), t.get('time', 0)) for t in data['tests']]
    times_only = [t[1] for t in test_times]

    if not times_only:
        return {
            "average_time": 0,
            "median_time": 0,
            "slowest_tests": [],
            "fastest_tests": [],
            "total_time": data['summary']['execution_time'],
            "insights": ["No test execution time data available"]
        }

    # Calculate statistics
    avg_time = sum(times_only) / len(times_only) if times_only else 0
    median_time = statistics.median(times_only) if times_only else 0

    # Sort tests by execution time
    sorted_tests = sorted(test_times, key=lambda x: x[1], reverse=True)
    slowest_tests = sorted_tests[:5]  # Top 5 slowest tests
    fastest_tests = sorted(test_times, key=lambda x: x[1])[:5]  # Top 5 fastest tests

    # Generate insights
    insights = []
    if len(times_only) > 1:
        if max(times_only) > avg_time * 3:
            insights.append("Some tests are significantly slower than average (>3x). Consider investigating these outliers.")

        if statistics.stdev(times_only) > avg_time:
            insights.append("High variation in test execution times. This might indicate inconsistent test complexity or environment issues.")

        # Look for patterns in failed tests
        failed_tests = [t for t in data['tests'] if t.get('status') == 'failed']
        failed_times = [t.get('time', 0) for t in failed_tests]
        if failed_times:
            avg_failed_time = sum(failed_times) / len(failed_times)
            if avg_failed_time > avg_time * 1.5:
                insights.append("Failed tests are taking longer than average. This might indicate performance-related failures.")
            elif avg_failed_time < avg_time * 0.5:
                insights.append("Failed tests are terminating faster than average. This might indicate early assertion failures.")

    return {
        "average_time": avg_time,
        "median_time": median_time,
        "slowest_tests": slowest_tests,
        "fastest_tests": fastest_tests,
        "total_time": data['summary']['execution_time'],
        "insights": insights if insights else ["No significant timing patterns detected"]
    }

def analyze_failure_patterns(data):
    """
    Analyze failed tests to identify common patterns and issues.
    Provides insights into the most common types of failures.
    """
    failures = [t for t in data['tests'] if t.get('status') == 'failed']
    if not failures:
        return {
            "common_errors": [],
            "affected_areas": [],
            "insights": ["No failures detected in this test run."]
        }

    # Extract error messages and types
    error_messages = [f.get('message', '') for f in failures if 'message' in f]
    error_types = [f.get('type', '') for f in failures if 'type' in f]

    # Count occurrences of each error type
    error_type_counts = {}
    for error_type in error_types:
        if error_type:
            error_type_counts[error_type] = error_type_counts.get(error_type, 0) + 1

    common_errors = sorted(
        [(k, v) for k, v in error_type_counts.items()],
        key=lambda x: x[1],
        reverse=True
    )[:5]  # Top 5 most common error types

    # Try to identify affected areas (e.g., from class names or features)
    affected_areas = {}
    for failure in failures:
        area = None
        # Try to extract area from classname or feature
        if 'classname' in failure:
            parts = failure['classname'].split('.')
            area = parts[0] if parts else None
        elif 'feature' in failure:
            area = failure['feature']

        if area:
            affected_areas[area] = affected_areas.get(area, 0) + 1

    top_affected_areas = sorted(
        [(k, v) for k, v in affected_areas.items()],
        key=lambda x: x[1],
        reverse=True
    )[:5]  # Top 5 most affected areas

    # Generate insights
    insights = []

    if common_errors:
        top_error_type, count = common_errors[0]
        if count > 1:
            insights.append(f"Most common error type: '{top_error_type}' ({count} occurrences)")

    if top_affected_areas:
        top_area, count = top_affected_areas[0]
        if count > 1:
            insights.append(f"Most affected area: '{top_area}' ({count} failures)")

    if AI_ANALYSIS_AVAILABLE and error_messages:
        # Use NLTK for more advanced text analysis
        try:
            # Combine all error messages
            combined_text = ' '.join(error_messages)

            # Tokenize and filter out stopwords
            stop_words = set(stopwords.words('english'))
            tokens = word_tokenize(combined_text.lower())
            filtered_tokens = [w for w in tokens if w.isalpha() and w not in stop_words]

            # Find most common terms in error messages
            from collections import Counter
            common_terms = Counter(filtered_tokens).most_common(5)

            if common_terms:
                terms_str = ", ".join([f"'{term}'" for term, count in common_terms])
                insights.append(f"Common terms in error messages: {terms_str}")
        except Exception:
            # Fallback if NLTK analysis fails
            pass

    if not insights:
        insights.append("No clear patterns found in the failed tests.")

    return {
        "common_errors": common_errors,
        "affected_areas": top_affected_areas,
        "insights": insights
    }

def generate_test_insights(data):
    """
    Generate comprehensive AI insights from test data.
    Combines multiple analysis functions to provide a complete picture.
    """
    # Gather all analysis data
    time_analysis = analyze_test_times(data)
    failure_analysis = analyze_failure_patterns(data)

    # Calculate additional metrics using the TestMetricsCalculator
    coverage_metrics = TestMetricsCalculator.calculate_coverage_metrics(data)
    flakiness_metrics = TestMetricsCalculator.calculate_flakiness_metrics(data)

    # Calculate overall metrics
    pass_rate = data['summary']['passed'] / data['summary']['total'] * 100 if data['summary']['total'] > 0 else 0

    # Generate insights about test health
    health_insights = []

    # Pass rate insights
    if pass_rate == 100:
        health_insights.append("All tests passed! Great job!")
    elif pass_rate >= 90:
        health_insights.append(f"High pass rate ({pass_rate:.1f}%). Look into the few failing tests to achieve 100%.")
    elif pass_rate >= 70:
        health_insights.append(f"Moderate pass rate ({pass_rate:.1f}%). Several tests need attention.")
    else:
        health_insights.append(f"Low pass rate ({pass_rate:.1f}%). Significant test failures to address.")

    # Test execution insights
    if data['summary']['execution_time'] > 300:  # More than 5 minutes
        health_insights.append(f"Test suite is taking a long time to execute ({data['summary']['execution_time']:.1f} seconds). Consider optimizing slow tests.")

    # Skipped tests insights
    skip_rate = data['summary']['skipped'] / data['summary']['total'] * 100 if data['summary']['total'] > 0 else 0
    if skip_rate > 20:
        health_insights.append(f"High number of skipped tests ({skip_rate:.1f}%). These tests might need maintenance.")

    # Flakiness insights
    if flakiness_metrics["has_flakiness_data"] and flakiness_metrics["flaky_tests"] > 0:
        flaky_rate = flakiness_metrics["flaky_rate"]
        health_insights.append(f"Detected {flakiness_metrics['flaky_tests']} flaky tests ({flaky_rate:.1f}%). Consider stabilizing these tests.")

    # Coverage insights
    if coverage_metrics["has_coverage_data"] or coverage_metrics["functional_coverage"] > 0:
        func_coverage = coverage_metrics["functional_coverage"]
        if func_coverage < 70:
            health_insights.append(f"Test coverage is relatively low ({func_coverage:.1f}%). Consider adding more tests.")
        elif func_coverage >= 90:
            health_insights.append(f"Great test coverage ({func_coverage:.1f}%). Keep up the good work!")

    # Historical comparison (placeholder for future enhancement)
    # This would require storing and retrieving previous test runs

    # Combine all insights
    all_insights = {
        "summary": {
            "pass_rate": pass_rate,
            "skip_rate": skip_rate,
            "execution_time": data['summary']['execution_time']
        },
        "timing": time_analysis,
        "failures": failure_analysis,
        "health_insights": health_insights,
        "coverage": coverage_metrics,
        "flakiness": flakiness_metrics,
        "key_metrics": {
            "reliability": min(pass_rate, 100),  # Percentage
            "performance": 100 - min(time_analysis["average_time"] / 2 * 100, 50),  # Simplified performance score
            "maintenance_burden": min(skip_rate + (data['summary']['failed'] / data['summary']['total'] * 100 if data['summary']['total'] > 0 else 0), 100),  # Higher is worse
            "test_coverage": coverage_metrics["functional_coverage"],  # Add coverage metric
            "flakiness_score": flakiness_metrics["flakiness_score"]  # Add flakiness score (higher is better)
        }
    }

    return all_insights

def create_advanced_visualizations(data, insights):
    """Create advanced visualization charts from test data and insights.

    Args:
        data: The parsed test data
        insights: The AI-generated insights about the test data

    Returns:
        BytesIO: A buffer containing the visualization image
    """
    if not VISUALIZATION_AVAILABLE:
        return None

    # Create a figure with multiple subplots
    fig = plt.figure(figsize=(12, 10))
    gs = plt.GridSpec(3, 2, figure=fig)

    # 1. Test Results Pie Chart (top left)
    ax1 = fig.add_subplot(gs[0, 0])
    results = [data['summary']['passed'], data['summary']['failed'], data['summary']['skipped']]
    labels = ['Passed', 'Failed', 'Skipped']
    colors = ['#4CAF50', '#F44336', '#FFC107']

    ax1.pie(results, labels=labels, autopct='%1.1f%%', startangle=90, colors=colors)
    ax1.set_title('Test Results')
    ax1.axis('equal')

    # 2. Key Metrics Radar Chart (top right)
    if 'key_metrics' in insights:
        ax2 = fig.add_subplot(gs[0, 1], polar=True)
        metrics = insights['key_metrics']

        # Define metrics for radar chart - invert maintenance_burden so lower is better
        categories = ['Reliability', 'Performance', 'Maintainability']
        values = [
            metrics['reliability'],
            metrics['performance'],
            100 - metrics['maintenance_burden']  # Invert so higher is better
        ]

        # Create radar chart
        angles = np.linspace(0, 2*np.pi, len(categories), endpoint=False).tolist()
        values = values + [values[0]]  # Close the loop
        angles = angles + [angles[0]]  # Close the loop
        categories = categories + [categories[0]]  # Close the loop

        ax2.plot(angles, values, 'o-', linewidth=2)
        ax2.fill(angles, values, alpha=0.25)
        ax2.set_thetagrids(np.degrees(angles[:-1]), categories[:-1])
        ax2.set_ylim(0, 100)
        ax2.set_title('Test Quality Metrics')
        ax2.grid(True)

    # 3. Test Execution Times Bar Chart (middle left)
    ax3 = fig.add_subplot(gs[1, 0])
    if len(data['tests']) > 0 and 'time' in data['tests'][0]:
        # Get top 5 slowest tests
        test_times = [(t.get('name', 'Unknown'), t.get('time', 0)) for t in data['tests']]
        sorted_tests = sorted(test_times, key=lambda x: x[1], reverse=True)[:5]

        test_names = [t[0][:20] + '...' if len(t[0]) > 20 else t[0] for t in sorted_tests]
        times = [t[1] for t in sorted_tests]

        bars = ax3.barh(test_names, times, color='#2196F3')
        ax3.set_xlabel('Execution Time (seconds)')
        ax3.set_title('Top 5 Slowest Tests')

        # Add times at the end of bars
        for i, (time, bar) in enumerate(zip(times, bars)):
            ax3.text(bar.get_width() + 0.1, bar.get_y() + bar.get_height()/2, f'{time:.2f}s',
                    ha='left', va='center')
    else:
        ax3.text(0.5, 0.5, 'No timing data available', ha='center', va='center')
        ax3.set_title('Test Execution Times')

    # 4. Error Distribution (middle right)
    ax4 = fig.add_subplot(gs[1, 1])
    if 'failures' in insights and insights['failures']['common_errors']:
        error_types = [e[0] if len(e[0]) < 20 else e[0][:17]+'...' for e in insights['failures']['common_errors']]
        error_counts = [e[1] for e in insights['failures']['common_errors']]

        bars = ax4.bar(error_types, error_counts, color='#F44336')
        ax4.set_xlabel('Error Type')
        ax4.set_ylabel('Count')
        ax4.set_title('Common Error Types')
        plt.xticks(rotation=45, ha='right')

        # Add counts on top of bars
        for i, count in enumerate(error_counts):
            ax4.text(i, count + 0.1, str(count), ha='center')
    else:
        ax4.text(0.5, 0.5, 'No errors detected', ha='center', va='center')
        ax4.set_title('Error Distribution')

    # 5. Test Execution Time Histogram (bottom)
    ax5 = fig.add_subplot(gs[2, :])
    if len(data['tests']) > 0 and 'time' in data['tests'][0]:
        times = [t.get('time', 0) for t in data['tests']]

        # Calculate histogram
        ax5.hist(times, bins=20, alpha=0.7, color='#009688')
        ax5.axvline(insights['timing']['average_time'], color='r', linestyle='dashed', linewidth=1, label=f'Mean: {insights["timing"]["average_time"]:.2f}s')
        ax5.axvline(insights['timing']['median_time'], color='g', linestyle='dashed', linewidth=1, label=f'Median: {insights["timing"]["median_time"]:.2f}s')

        ax5.set_xlabel('Execution Time (seconds)')
        ax5.set_ylabel('Frequency')
        ax5.set_title('Test Execution Time Distribution')
        ax5.legend()
    else:
        ax5.text(0.5, 0.5, 'No timing data available', ha='center', va='center')
        ax5.set_title('Test Execution Time Distribution')

    plt.tight_layout()

    # Save to BytesIO
    buffer = BytesIO()
    plt.savefig(buffer, format='png', dpi=100)
    buffer.seek(0)
    plt.close(fig)

    return buffer

def format_duration(seconds):
    """Format seconds into a human-readable duration format.

    Args:
        seconds (float): Duration in seconds

    Returns:
        str: Formatted duration string (e.g., "2m 30s", "1h 15m 30s")
    """
    if seconds is None or seconds < 0:
        return "0s"

    # Convert to float to ensure proper handling of decimal values
    try:
        seconds = float(seconds)
    except (ValueError, TypeError):
        return "0s"

    # Handle zero or very small values
    if seconds == 0:
        return "0s"
    elif seconds < 0.0001:  # For extremely small values
        return "< 0.0001s"

    hours, remainder = divmod(seconds, 3600)
    minutes, seconds = divmod(remainder, 60)

    if hours > 0:
        return f"{int(hours)}h {int(minutes)}m {seconds:.2f}s"
    elif minutes > 0:
        return f"{int(minutes)}m {seconds:.2f}s"
    else:
        # Handle milliseconds for short durations
        if seconds < 0.1:
            return f"{seconds:.4f}s"  # Show 4 decimal places for very small durations
        elif seconds < 1:
            return f"{seconds:.3f}s"  # Show 3 decimal places for durations under 1s
        else:
            return f"{seconds:.2f}s"  # Show 2 decimal places for all other durations

def generate_html_report(data, project_name, include_insights=True, include_charts=True):  # Added include_charts parameter
    """Generate an HTML report from the parsed test data with AI insights."""
    # Generate insights if enabled
    insights = generate_test_insights(data) if include_insights else None
    visualization_buffer = None
    if include_charts and VISUALIZATION_AVAILABLE and insights:  # Check for insights as well, as charts use them
        visualization_buffer = create_advanced_visualizations(data, insights)

    # Format the total execution time
    total_execution_time = data['summary']['execution_time']
    formatted_execution_time = format_duration(total_execution_time)

    # Prepare HTML template
    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{project_name} - Test Documentation</title>
    <style>
        body {{
            font-family: Arial, sans-serif;
            line-height: 1.6;
            max-width: 1200px;
            margin: 0 auto;
            padding: 20px;
            color: #333;
            background-color: #f4f7f6; /* Light background for the page */
        }}
        header {{
            background-color: #0D47A1; /* Dark Blue */
            color: white;
            padding: 20px;
            margin-bottom: 20px;
            border-radius: 8px;
            text-align: center;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        }}
        header h1 {{
            margin: 0;
            color: white; /* Ensure h1 in header is white */
        }}
        header p {{
            margin-top: 5px;
            font-size: 0.9em;
            opacity: 0.8;
        }}
        h1, h2, h3 {{
            color: #0D47A1; /* Dark Blue for general headings */
        }}
        h2 {{
            border-bottom: 2px solid #B0BEC5; /* Light Grey Blue */
            padding-bottom: 10px;
            margin-top: 40px; /* More space above section titles */
            margin-bottom: 20px;
        }}
        .summary {{
            display: grid;
            grid-template-columns: repeat(auto-fill, minmax(200px, 1fr));
            gap: 20px; /* Increased gap */
            margin-bottom: 30px;
        }}
        .metric {{
            background-color: #FFFFFF; /* White background for metrics */
            border-left: 5px solid #0D47A1; /* Dark Blue */
            padding: 20px; /* Increased padding */
            border-radius: 0 8px 8px 0; /* Rounded corners */
            box-shadow: 0 2px 5px rgba(0,0,0,0.05);
        }}
        .metric h3 {{
            margin-top: 0;
            margin-bottom: 10px; /* Space below metric title */
            color: #1565C0; /* Medium Blue */
            font-size: 1.1em;
        }}
        .metric p {{
            font-size: 1.5em; /* Larger metric value */
            font-weight: bold;
            margin-bottom: 0;
        }}
        table {{
            width: 100%;
            border-collapse: collapse;
            margin-bottom: 30px;
            box-shadow: 0 2px 8px rgba(0,0,0,0.1); /* Softer shadow */
            border-radius: 8px; /* Rounded corners for table */
            overflow: hidden; /* Ensures border-radius clips content */
        }}
        th, td {{
            padding: 12px 15px;
            text-align: left;
            border-bottom: 1px solid #CFD8DC; /* Grey Blue */
        }}
        th {{
            background-color: #1565C0; /* Medium Blue */
            color: white;
            font-weight: bold; /* Bold header text */
        }}
        tr:nth-child(even) {{
            background-color: #F8F9FA; /* Very Light Grey for even rows */
        }}
        tr:hover {{
            background-color: #E3F2FD; /* Light Blue hover */
        }}
        .passed {{
            color: #2E7D32; /* Dark Green */
            font-weight: bold;
        }}
        .failed {{
            color: #C62828; /* Dark Red */
            font-weight: bold;
        }}
        .skipped {{
            color: #FF8F00; /* Dark Orange */
            font-weight: bold;
        }}
        footer {{
            margin-top: 50px;
            text-align: center;
            font-size: 0.9em; /* Slightly larger footer text */
            color: #78909C; /* Grey Blue */
            padding: 20px 0;
            border-top: 1px solid #CFD8DC;
        }}
        .card {{
            background-color: white;
            border-radius: 8px;
            box-shadow: 0 4px 8px rgba(0,0,0,0.08); /* Slightly more pronounced shadow */
            padding: 25px; /* Increased padding for cards */
            margin-bottom: 30px; /* Increased margin for cards */
        }}
        .insights-section h2 {{ /* Specific styling for AI Insights title */
             margin-bottom: 25px;
        }}
        .section-title {{ /* For AI Insights sub-section titles (h3) */
            margin-top: 25px;
            margin-bottom: 15px;
            color: #1565C0; /* Medium Blue */
            font-size: 1.3em; /* Slightly larger sub-section titles */
            border-bottom: 1px solid #E0E0E0;
            padding-bottom: 8px;
        }}
        ul.insights-list {{
            list-style-type: none;
            padding-left: 0;
        }}
        ul.insights-list li {{
            background-color: #E1F5FE; /* Very Light Blue */
            padding: 12px 15px; /* Adjusted padding */
            margin-bottom: 8px;
            border-left: 4px solid #0288D1; /* Medium-Dark Blue */
            border-radius: 4px;
            font-size: 0.95em;
        }}
        ul.insights-list h4 {{ /* For 'Slowest Tests'/'Most Affected Areas' titles */
            margin-top: 15px;
            margin-bottom: 10px;
            color: #0D47A1;
            font-size: 1em;
        }}
        .key-metrics-container {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(280px, 1fr)); /* Adjusted minmax for better fit */
            gap: 20px;
            margin-bottom: 30px;
        }}
        .key-metric-card {{
            padding: 25px; /* Increased padding */
            border-radius: 8px;
            color: white;
            text-align: center;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
            display: flex;
            flex-direction: column;
            justify-content: center;
        }}
        .reliability {{
            background: linear-gradient(135deg, #4CAF50, #388E3C); /* Green shades */
        }}
        .performance {{
            background: linear-gradient(135deg, #2196F3, #1976D2); /* Blue shades */
        }}
        .maintenance {{
            background: linear-gradient(135deg, #FFC107, #FFA000); /* Amber/Yellow shades */
        }}
        .metric-title {{
            font-size: 1.2em;
            font-weight: bold;
            margin-bottom: 10px;
        }}
        .metric-value {{
            font-size: 2.8em;
            font-weight: bold;
            margin: 8px 0; /* Adjusted margin */
        }}
        .metric-description {{
            font-size: 0.9em;
            opacity: 0.9;
        }}
        .visualization-container {{
            text-align: center;
            margin: 30px 0;
            padding: 20px;
            background-color: #FFFFFF; /* White background */
            border-radius: 8px;
            /* box-shadow: 0 2px 5px rgba(0,0,0,0.1); Already a card */
        }}
        .visualization-container h2 {{
             margin-bottom: 20px;
        }}
        .visualization-container img {{
            max-width: 100%;
            height: auto;
            border-radius: 5px;
            border: 1px solid #CFD8DC; /* Light border for the image */
        }}
        .failure-details .failure {{ /* Already a card, this is for individual failure blocks */
            background-color: #FFEBEE; /* Light Pink */
            border: 1px solid #E57373; /* Pink */
            border-left: 5px solid #C62828; /* Dark Red */
            padding: 20px; /* Increased padding */
            margin-bottom: 20px; /* Increased margin */
            border-radius: 5px;
        }}
        .failure-details h3 {{ /* For individual failure titles */
            color: #C62828; /* Dark Red */
            margin-top:0;
            margin-bottom: 10px;
            font-size: 1.1em;
        }}
        .failure-details pre {{
            background-color: #FCE4EC; /* Lighter Pink */
            padding: 15px; /* Increased padding */
            border-radius: 4px;
            overflow-x: auto;
            border: 1px solid #F8BBD0; /* Light Pink Border */
            font-size: 0.85em; /* Smaller font for stacktrace */
            line-height: 1.4;
        }}
        .metrics-grid {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(250px, 1fr)); /* Consistent card sizing */
            gap: 20px;
            margin: 20px 0;
        }}
        .metric-card {{
            background-color: #FFFFFF;
            border-radius: 12px;
            box-shadow: 0 4px 12px rgba(0,0,0,0.15);
            overflow: hidden;
            display: flex;
            flex-direction: column;
            transition: transform 0.2s, box-shadow 0.2s;
            height: 100%;
        }}
        .metric-card:hover {{
            transform: translateY(-5px);
            box-shadow: 0 8px 16px rgba(0,0,0,0.2);
        }}
        .metric-header {{
            padding: 15px;
            color: white;
            display: flex;
            align-items: center;
            justify-content: space-between;
        }}
        .metric-icon {{
            font-size: 2.2em;
            margin-right: 10px;
            display: flex;
            align-items: center;
            justify-content: center;
            height: 50px;
            width: 50px;
            background-color: rgba(255,255,255,0.2);
            border-radius: 50%;
        }}
        .metric-body {{
            padding: 20px;
            text-align: center;
            display: flex;
            flex-direction: column;
            justify-content: center;
            align-items: center;
            flex-grow: 1;
        }}
        .metric-title {{
            font-size: 1.2em;
            font-weight: 600;
            margin-bottom: 10px;
            color: #333;
        }}
        .metric-value {{
            font-size: 3em;
            font-weight: 700;
            margin: 10px 0;
            line-height: 1;
        }}
        .metric-desc {{
            font-size: 0.9em;
            color: #757575;
            text-align: center;
            margin-top: 10px;
        }}
        .gauge-container {{
            width: 100%;
            margin: 12px 0;
        }}
        .metric-gauge {{
            height: 10px;
            background-color: rgba(0,0,0,0.1);
            border-radius: 5px;
            overflow: hidden;
        }}
        .gauge-fill {{
            height: 100%;
            border-radius: 5px;
            transition: width 1s ease-out;
        }}
        /* Media queries for responsiveness */
        @media (max-width: 768px) {{
            .metrics-grid {{
                grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
            }}
        }}
        @media (max-width: 480px) {{
            .metrics-grid {{
                grid-template-columns: 1fr;
            }}
        }}
    </style>
</head>
<body>
    <header>
        <h1>{project_name} - Test Documentation</h1>
        <p>Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
    </header>

    <main>
        <section class="card">
            <h2>Test Summary</h2>
            <div class="summary">
                <div class="metric">
                    <h3>Total Tests</h3>
                    <p>{data['summary']['total']}</p>
                </div>
                <div class="metric">
                    <h3>Passed</h3>
                    <p class="passed">{data['summary']['passed']}</p>
                </div>
                <div class="metric">
                    <h3>Failed</h3>
                    <p class="failed">{data['summary']['failed']}</p>
                </div>
                <div class="metric">
                    <h3>Skipped</h3>
                    <p class="skipped">{data['summary']['skipped']}</p>
                </div>
                <div class="metric">
                    <h3>Execution Time</h3>
                    <p>{formatted_execution_time}</p>
                </div>
            </div>
        </section>
    """
    # AI Insights Section
    if include_insights and insights:
        html += """
        <section class="card insights-section">
            <h2>AI-Powered Test Insights</h2>
        """

        # Key Quality Metrics - Enhanced visual representation with properly closed tags
        if 'key_metrics' in insights:
            html += """
            <div class="key-metrics-container">
                <h3 class="section-title">Key Quality Metrics</h3>
                <div class="metrics-grid">
"""
            # Reliability metric
            reliability = insights['key_metrics']['reliability']
            reliability_color = "#4CAF50" if reliability >= 90 else ("#FFC107" if reliability >= 70 else "#F44336")
            html += f"""
                    <div style="flex: 1; min-width: 200px; background: linear-gradient(135deg, {reliability_color}, {reliability_color}80); color: white; padding: 20px; border-radius: 8px; text-align: center;">
                        <h4 style="margin-top: 0;">Reliability Score</h4>
                        <div style="font-size: 2.5em; font-weight: bold; margin: 10px 0;">{reliability:.1f}%</div>
                        <p style="margin-bottom: 0; font-size: 0.9em;">Based on test pass rate</p>
                    </div>
                """

            # Performance metric
            performance = insights['key_metrics']['performance']
            performance_color = "#2196F3" if performance >= 90 else "#03A9F4" if performance >= 70 else "#B3E5FC"
            html += f"""
                    <div style="flex: 1; min-width: 200px; background: linear-gradient(135deg, {performance_color}, {performance_color}80); color: white; padding: 20px; border-radius: 8px; text-align: center;">
                        <h4 style="margin-top: 0;">Performance Score</h4>
                        <div style="font-size: 2.5em; font-weight: bold; margin: 10px 0;">{performance:.1f}%</div>
                        <p style="margin-bottom: 0; font-size: 0.9em;">Based on execution time efficiency</p>
                    </div>
                """

            # Maintenance burden metric
            burden = insights['key_metrics']['maintenance_burden']
            # Reverse the scale for maintenance burden (lower is better)
            burden_score = 100 - burden
            burden_color = "#4CAF50" if burden <= 10 else "#FFC107" if burden <= 30 else "#F44336"
            html += f"""
                <div style="flex: 1; min-width: 200px; background: linear-gradient(135deg, {burden_color}, {burden_color}80); color: white; padding: 20px; border-radius: 8px; text-align: center;">
                    <h4 style="margin-top: 0;">Maintainability Score</h4>
                    <div style="font-size: 2.5em; font-weight: bold; margin: 10px 0;">{burden_score:.1f}%</div>
                    <p style="margin-bottom: 0; font-size: 0.9em;">Based on failures and skipped tests</p>
                </div>
            """

            # Test coverage metric
            if 'test_coverage' in insights['key_metrics']:
                coverage = insights['key_metrics']['test_coverage']
                coverage_color = "#4CAF50" if coverage >= 80 else "#FFC107" if coverage >= 60 else "#F44336"
                html += f"""
                    <div style="flex: 1; min-width: 200px; background: linear-gradient(135deg, {coverage_color}, {coverage_color}80); color: white; padding: 20px; border-radius: 8px; text-align: center;">
                        <h4 style="margin-top: 0;">Coverage Score</h4>
                        <div style="font-size: 2.5em; font-weight: bold; margin: 10px 0;">{coverage:.1f}%</div>
                        <p style="margin-bottom: 0; font-size: 0.9em;">Based on functional test coverage</p>
                    </div>
                """

            # Flakiness score metric
            if 'flakiness_score' in insights['key_metrics']:
                flakiness = insights['key_metrics']['flakiness_score']
                flakiness_color = "#4CAF50" if flakiness >= 90 else "#FFC107" if flakiness >= 70 else "#F44336"
                html += f"""
                    <div style="flex: 1; min-width: 200px; background: linear-gradient(135deg, {flakiness_color}, {flakiness_color}80); color: white; padding: 20px; border-radius: 8px; text-align: center;">
                        <h4 style="margin-top: 0;">Stability Score</h4>
                        <div style="font-size: 2.5em; font-weight: bold; margin: 10px 0;">{flakiness:.1f}%</div>
                        <p style="margin-bottom: 0; font-size: 0.9em;">Based on test reliability</p>
                    </div>
                """

            html += """
                </div>
            </div>
            """

        # Health insights
        if 'health_insights' in insights and insights['health_insights']:
            html += """
            <div class="insights-block">
                <h3>Health Assessment</h3>
                <ul class="insights-list">
            """
            for insight in insights['health_insights']:
                html += f"<li>{insight}</li>"
            html += """
                </ul>
            </div>
            """

        # Performance insights
        if 'timing' in insights and insights['timing']['insights']:
            html += """
            <div class="insights-block">
                <h3>Performance Insights</h3>
                <ul class="insights-list">
            """
            for insight in insights['timing']['insights']:
                html += f"<li>{insight}</li>"
            html += """
                </ul>
            """

            # Show slowest tests
            if insights['timing']['slowest_tests']:
                html += """
                <h4>Slowest Tests</h4>
                <ol>
                """
                for name, time in insights['timing']['slowest_tests'][:3]:
                    html += f"<li>{name} - {format_duration(time)}</li>"
                html += """
                </ol>
                """

            html += "</div>"

        # Only display test failures if there are actual failures
        if 'failures' in insights and insights['failures'] and data['summary']['failed'] > 0:
            html += """
            <div class="insights-block">
                <h3>Test Failures Analysis</h3>
            """

            # Failure insights
            if 'insights' in insights['failures']:
                html += "<ul class='insights-list'>"
                for insight in insights['failures']['insights']:
                    if "No failures detected" not in insight:
                        html += f"<li>{insight}</li>"
                html += "</ul>"

            # Affected areas
            if 'affected_areas' in insights['failures'] and insights['failures']['affected_areas']:
                html += "<h4>Most Affected Areas</h4><ul>"
                for area, count in insights['failures']['affected_areas']:
                    html += f"<li>{area} - {count} failures</li>"
                html += "</ul>"

            # Display actual failures
            html += "<h3>Failure Details</h3>"
            failures = [t for t in data['tests'] if t.get('status') == 'failed']
            if failures:
                for i, failure in enumerate(failures):
                    html += f"""
                    <div class="failure">
                        <h3>{i+1}. {failure.get('name', 'Unknown Test')}</h3>
                        <p><strong>Message:</strong> {failure.get('message', 'No message available')}</p>
                        <p><strong>Type:</strong> {failure.get('type', 'Unknown error type')}</p>
                        <pre>{failure.get('stacktrace', 'No stacktrace available')}</pre>
                    </div>
                """

            html += "</div>"
        elif data['summary']['failed'] == 0:
            html += """
            <div class="insights-block success-block" style="background-color: #E8F5E9; padding: 15px; border-left: 5px solid #4CAF50; border-radius: 4px;">
                <h3 style="color: #2E7D32; margin-top: 0;">No Failures Detected</h3>
                <p>All tests passed successfully. Great job!</p>
            </div>
            """

        # Add a new summary section that provides an overall assessment
        html += """
        <div class="insights-block summary-section">
            <h3>Executive Summary</h3>
            <div style="background-color: #F5F5F5; padding: 20px; border-radius: 8px; margin-bottom: 20px;">
        """

        # Generate smart summary based on test results
        total_tests = data['summary']['total']
        pass_rate = insights['summary']['pass_rate']
        execution_time = data['summary']['execution_time']

        if pass_rate == 100:
            summary_class = "success"
            emoji = "✅"
            summary_title = "Perfect Test Run"
            summary_text = f"All {total_tests} tests passed successfully in {formatted_execution_time}."
        elif pass_rate >= 90:
            summary_class = "near-success"
            emoji = "🎯"
            summary_title = "High Quality Test Run"
            failed = data['summary']['failed']
            summary_text = f"Test suite is performing well with {pass_rate:.1f}% pass rate. Only {failed} test(s) failed."
        elif pass_rate >= 70:
            summary_class = "caution"
            emoji = "⚠️"
            summary_title = "Moderate Quality Test Run"
            failed = data['summary']['failed']
            summary_text = f"Test quality needs improvement. {failed} tests failed ({pass_rate:.1f}% pass rate)."
        else:
            summary_class = "danger"
            emoji = "❌"
            summary_title = "Low Quality Test Run"
            failed = data['summary']['failed']
            summary_text = f"Critical issues detected with {failed} failed tests. Only {pass_rate:.1f}% of tests passed."

        html += f"""
            <h2 style="margin-top: 0;">{emoji} {summary_title}</h2>
            <p style="font-size: 1.2em;">{summary_text}</p>
            <div style="display: flex; flex-wrap: wrap; gap: 10px; margin-top: 15px;">
        """

        # Add key recommendations based on insights
        if pass_rate < 100:
            if 'timing' in insights and any("significantly slower" in insight for insight in insights['timing']['insights']):
                html += """<div style="background-color: #FFF8E1; padding: 10px; border-radius: 4px; margin-right: 10px;">
                    <strong>💡 Recommendation:</strong> Investigate and optimize slow-performing tests.
                </div>"""

        if 'skipped' in data['summary'] and data['summary']['skipped'] > 0:
            html += """<div style="background-color: #FFF8E1; padding: 10px; border-radius: 4px;">
                <strong>💡 Recommendation:</strong> Address skipped tests to improve test coverage.
            </div>"""

        if execution_time > 300:  # More than 5 minutes
            html += """<div style="background-color: #FFF8E1; padding: 10px; border-radius: 4px;">
                <strong>💡 Recommendation:</strong> Consider test optimization or parallelization to reduce execution time.
            </div>"""

        html += """
            </div>
        </div>
        """

        html += """
        </div>
    </section>
        """
    else:
        html += """
        <section class="card insights-section">
            <h2>AI-Powered Test Insights</h2>
            <p>No insights available.</p>
        </section>
        """

    html += """
        <section class="card">
            <h2>Test Details</h2>
            <table>
                <thead>
                    <tr>
                        <th>#</th>
                        <th>Test Name</th>
                        <th>Status</th>
                        <th>Duration (s)</th>
                    </tr>
                </thead>
                <tbody>
    """

    # Add test rows
    for i, test in enumerate(data['tests']):
        status_class = test.get('status', '').lower()
        test_time = test.get('time', 0)
        formatted_test_time = format_duration(test_time)

        html += f"""
                    <tr>
                        <td>{i+1}</td>
                        <td>{test.get('name', 'Unknown Test')}</td>
                        <td class="{status_class}">{test.get('status', 'unknown').upper()}</td>
                        <td>{formatted_test_time}</td>
                    </tr>"""

    html += """
                </tbody>
            </table>
        </section>
    """

    html += """
    </main>

    <footer>
        <p>Generated by Jarvis Test Automation Platform © 2025</p>
    </footer>
</body>
</html>
"""

    return html

def generate_markdown_report(data, project_name, include_insights=True):
    """Generate a Markdown report from the parsed test data with AI insights."""
    # Generate insights if enabled
    insights = generate_test_insights(data) if include_insights else None

    md = f"""# {project_name} - Test Documentation

Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

## Test Summary

- **Total Tests:** {data['summary']['total']}
- **Passed:** {data['summary']['passed']}
- **Failed:** {data['summary']['failed']}
- **Skipped:** {data['summary']['skipped']}
- **Execution Time:** {data['summary']['execution_time']:.2f} seconds
"""

    # Add key metrics and insights if available
    if insights:
        md += f"""
## Key Quality Metrics

| Metric | Score | Description |
|--------|-------|-------------|
| 🔵 **Reliability** | {insights['key_metrics']['reliability']:.1f}% | Based on test pass rate |
| 🟢 **Performance** | {insights['key_metrics']['performance']:.1f}% | Based on execution time efficiency |
| 🟠 **Maintenance Burden** | {insights['key_metrics']['maintenance_burden']:.1f}% | Based on failures and skipped tests |"""

        # Add coverage metric if available
        if 'test_coverage' in insights['key_metrics']:
            md += f"\n| 📊 **Test Coverage** | {insights['key_metrics']['test_coverage']:.1f}% | Based on functional test coverage |"

        # Add flakiness score if available
        if 'flakiness_score' in insights['key_metrics']:
            md += f"\n| 🧪 **Test Stability** | {insights['key_metrics']['flakiness_score']:.1f}% | Based on test reliability |"

        md += """

## AI-Powered Test Insights

### Health Assessment
"""

        # Add health insights
        for insight in insights['health_insights']:
            md += f"- {insight}\n"

        # Add performance insights
        if insights['timing']['insights']:
            md += "\n### Performance Insights\n"
            for insight in insights['timing']['insights']:
                md += f"- {insight}\n"

            # Add info about the slowest tests
            if insights['timing']['slowest_tests']:
                md += "\n**Slowest Tests:**\n"
                for i, (name, time) in enumerate(insights['timing']['slowest_tests'][:3]):
                    md += f"{i+1}. {name} - {time:.2f}s\n"

        # Add failure pattern insights
        if insights['failures']['insights']:
            md += "\n### Failure Analysis\n"
            for insight in insights['failures']['insights']:
                md += f"- {insight}\n"

            # Add affected areas if available
            if insights['failures']['affected_areas']:
                md += "\n**Most Affected Areas:**\n"
                for area, count in insights['failures']['affected_areas']:
                    md += f"- {area} - {count} failures\n"

    # Add test details
    md += "\n## Test Details\n\n"
    md += "| # | Test Name | Status | Duration |\n"
    md += "|---|-----------|--------|----------|\n"

    # Add test rows
    for i, test in enumerate(data['tests']):
        status = test.get('status', 'unknown').upper()
        test_time = test.get('time', 0)

        md += f"| {i+1} | {test.get('name', 'Unknown Test')} | {status} | {test_time:.2f}s |\n"

    # Add failure details
    md += "\n## Failure Details\n\n"

    failures = [t for t in data['tests'] if t.get('status') == 'failed']
    if failures:
        for i, failure in enumerate(failures):
            md += f"""### {i+1}. {failure.get('name', 'Unknown Test')}

**Message:** {failure.get('message', 'No message available')}

**Type:** {failure.get('type', 'Unknown error type')}

```
{failure.get('stacktrace', 'No stacktrace available')}
```

"""
    else:
        md += "No failures reported.\n"

    md += "\n---\nGenerated by Jarvis Test Automation Platform"

    return md

def generate_pdf_report(data, project_name):
    """Generate a PDF report from the parsed test data."""
    if not PDF_GENERATION_AVAILABLE:
        st.warning("PDF generation requires ReportLab. Install with: pip install reportlab")
        return None

    buffer = BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter

    # Add title and header
    p.setFont("Helvetica-Bold", 18)
    p.drawString(72, height - 72, f"{project_name} - Test Documentation")
    p.setFont("Helvetica", 12)
    p.drawString(72, height - 96, f"Generated on {datetime.now().strftime('%Y%m-%d %H:%M:%S')}")

    # Add summary information
    p.setFont("Helvetica-Bold", 14)
    p.drawString(72, height - 144, "Test Summary")

    p.setFont("Helvetica", 12)
    p.drawString(72, height - 168, f"Total Tests: {data['summary']['total']}")
    p.drawString(72, height - 184, f"Passed: {data['summary']['passed']}")
    p.drawString(72, height - 200, f"Failed: {data['summary']['failed']}")
    p.drawString(72, height - 216, f"Skipped: {data['summary']['skipped']}")
    p.drawString(72, height - 232, f"Execution Time: {data['summary']['execution_time']:.2f} seconds")

    # Add test details (simplified for PDF)
    p.setFont("Helvetica-Bold", 14)
    p.drawString(72, height - 280, "Test Results")

    y = height - 304
    p.setFont("Helvetica-Bold", 10)
    p.drawString(72, y, "Test Name")
    p.drawString(300, y, "Status")
    p.drawString(400, y, "Duration")

    p.line(72, y - 6, 520, y - 6)
    y -= 20

    # Add test rows (limited to fit on page)
    p.setFont("Helvetica", 10)

    max_tests = min(len(data['tests']), 20)  # Limit to first 20 tests
    for i in range(max_tests):
        test = data['tests'][i]
        status = test.get('status', 'unknown').upper()
        test_time = test.get('time', 0)

        # Check if we need a new page
        if y < 72:
            p.showPage()
            p.setFont("Helvetica-Bold", 14)
            p.drawString(72, height - 72, "Test Results (continued)")
            p.setFont("Helvetica", 10)
            y = height - 100

        p.drawString(72, y, test.get('name', 'Unknown Test')[:30])
        p.drawString(300, y, status)
        p.drawString(400, y, f"{test_time:.2f}s")

        y -= 16

    # Add footer
    p.setFont("Helvetica-Italic", 8)
    p.drawString(72, 40, "Generated by Jarvis Test Automation Platform")

    p.save()
    buffer.seek(0)
    return buffer.read()

def generate_docx_report(data, project_name):
    """Generate a Word document report from the parsed test data."""
    if not DOCX_GENERATION_AVAILABLE:
        st.warning("Word document generation requires python-docx. Install with: pip install python-docx")
        return None

    doc = Document()

    # Add title
    doc.add_heading(f"{project_name} - Test Documentation", level=1)
    doc.add_paragraph(f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    # Add summary information
    doc.add_heading("Test Summary", level=2)
    summary_table = doc.add_table(rows=5, cols=2)
    summary_table.style = 'Table Grid'

    # Fill summary table
    summary_items = [
        ("Total Tests", str(data['summary']['total'])),
        ("Passed", str(data['summary']['passed'])),
        ("Failed", str(data['summary']['failed'])),
        ("Skipped", str(data['summary']['skipped'])),
        ("Execution Time", f"{data['summary']['execution_time']:.2f} seconds")
    ]

    for i, (label, value) in enumerate(summary_items):
        row = summary_table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value

    # Add test details
    doc.add_heading("Test Details", level=2)
    test_table = doc.add_table(rows=1, cols=4)
    test_table.style = 'Table Grid'

    # Add table header
    header = test_table.rows[0]
    header.cells[0].text = "#"
    header.cells[1].text = "Test Name"
    header.cells[2].text = "Status"
    header.cells[3].text = "Duration"

    # Add test rows
    for i, test in enumerate(data['tests']):
        row = test_table.add_row()
        row.cells[0].text = str(i+1)
        row.cells[1].text = test.get('name', 'Unknown Test')
        row.cells[2].text = test.get('status', 'unknown').upper()
        row.cells[3].text = f"{test.get('time', 0):.2f}s"

    # Add failure details
    doc.add_heading("Failure Details", level=2)

    failures = [t for t in data['tests'] if t.get('status') == 'failed']
    if failures:
        for i, failure in enumerate(failures):
            doc.add_heading(f"{i+1}. {failure.get('name', 'Unknown Test')}", level=3)
            doc.add_paragraph(f"Message: {failure.get('message', 'No message available')}")
            doc.add_paragraph(f"Type: {failure.get('type', 'Unknown error type')}")
            doc.add_paragraph(failure.get('stacktrace', 'No stacktrace available'))
    else:
        doc.add_paragraph("No failures reported.")

    # Add footer
    doc.add_paragraph("Generated by Jarvis Test Automation Platform").italic = True

    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()

def detect_file_type(uploaded_file):
    """Detect the type of test result file based on content."""
    content = uploaded_file.read()
    uploaded_file.seek(0)  # Reset file pointer

    # Try to detect file type based on content
    content_str = content.decode('utf-8', errors='ignore')

    if re.search(r'<\?xml.*\?>', content_str[:100]):
        if "<robot" in content_str[:1000]:
            return "robot", content_str
        else:
            return "junit", content_str
    elif content_str.strip().startswith('{') or content_str.strip().startswith('['):
        try:
            # Try parsing as JSON
            json.loads(content_str)
            if '"elements":' in content_str and '"type":"scenario"' in content_str:
                return "cucumber", content_str
            else:
                return "json", content_str
        except:
            pass

    # Default to unknown
    return "unknown", content_str

def process_file(uploaded_file):
    """Process the uploaded test result file and extract data."""
    file_type, content_str = detect_file_type(uploaded_file)

    if file_type == "junit":
        return parse_junit_xml(content_str)
    elif file_type == "robot":
        return parse_robot_xml(content_str)
    elif file_type == "cucumber":
        return parse_cucumber_json(content_str)
    else:
        st.error(f"Unsupported file format for {uploaded_file.name}. Please upload JUnit XML, Robot Framework XML, or Cucumber JSON.")
        return None

def create_visualization(data):
    """Create visualization charts from test data."""
    if not VISUALIZATION_AVAILABLE:
        st.warning("Visualization requires matplotlib and seaborn. Install with: pip install matplotlib seaborn")
        return None

    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(12, 5))

    # Chart 1: Test results pie chart
    results = [data['summary']['passed'], data['summary']['failed'], data['summary']['skipped']]
    labels = ['Passed', 'Failed', 'Skipped']
    colors = ['#4CAF50', '#F44336', '#FFC107']

    ax1.pie(results, labels=labels, autopct='%1.1f%%', startangle=90, colors=colors)
    ax1.set_title('Test Results')
    ax1.axis('equal')

    # Chart 2: Test execution times
    if len(data['tests']) > 0 and 'time' in data['tests'][0]:
        test_names = [t.get('name', 'Unknown')[:20] for t in data['tests'][:10]]  # First 10 tests
        times = [t.get('time', 0) for t in data['tests'][:10]]
        statuses = [t.get('status', 'unknown') for t in data['tests'][:10]]

        # Map statuses to colors
        status_colors = {'passed': '#4CAF50', 'failed': '#F44336', 'skipped': '#FFC107', 'unknown': '#9E9E9E'}
        bar_colors = [status_colors.get(s.lower(), '#9E9E9E') for s in statuses]

        ax2.barh(test_names, times, color=bar_colors)
        ax2.set_xlabel('Execution Time (seconds)')
        ax2.set_title('Test Execution Times (Top 10)')

    plt.tight_layout()

    # Save to BytesIO
    buffer = BytesIO()
    plt.savefig(buffer, format='png')
    buffer.seek(0)
    return buffer

def show_ui():
    """Main function to display the UI for automatic test documentation generation."""
    st.title("Automated Test Documentation")
    st.markdown("""
    This tool automatically generates comprehensive documentation from your test results.
    Upload test result files (JUnit XML, Robot Framework XML, or Cucumber JSON) to get started.
    """)

    # Create tabs
    upload_tab, config_tab, preview_tab = st.tabs(["Upload Results", "Configuration", "Preview"])

    with upload_tab:
        st.subheader("Upload Test Results")

        uploaded_files = st.file_uploader("Upload test result files",
                                         type=["xml", "json"],
                                         accept_multiple_files=True,
                                         help="Upload JUnit XML, Robot Framework XML, or Cucumber JSON files")

        if uploaded_files:
            st.success(f"Uploaded {len(uploaded_files)} file(s)")

            # Process files
            results_data = {}
            for uploaded_file in uploaded_files:
                with st.spinner(f"Processing {uploaded_file.name}"):
                    # Process the file
                    file_data = process_file(uploaded_file)

                    if file_data:
                        results_data[uploaded_file.name] = file_data
                        st.success(f"✅ Successfully processed {uploaded_file.name}")
                    else:
                        st.error(f"❌ Failed to process {uploaded_file.name}")

            if results_data:
                st.session_state.test_results = results_data
                st.session_state.current_file = next(iter(results_data.keys()))

    with config_tab:
        st.subheader("Documentation Configuration")

        col1, col2 = st.columns(2)
        with col1:
            project_name = st.text_input("Project Name", "Test Project",
                                        help="Name of the project for documentation")

        with col2:
            doc_format = st.selectbox("Documentation Format",
                                     list(OUTPUT_FORMATS.keys()),
                                     help="Format of the generated documentation")

        st.subheader("Content Options")

        col1, col2, col3 = st.columns(3)
        with col1:
            include_summary = st.checkbox("Include Summary", value=True,
                                         help="Include test execution summary")
        with col2:
            include_details = st.checkbox("Include Test Details", value=True,
                                         help="Include detailed test results")
        with col3:
            include_failures = st.checkbox("Include Failure Details", value=True,
                                          help="Include detailed information about test failures")

        if VISUALIZATION_AVAILABLE:
            include_charts = st.checkbox("Include Charts", value=True,
                                        help="Include visual charts in the documentation")
        else:
            include_charts = False
            st.info("Charts require matplotlib and seaborn. Install with: pip install matplotlib seaborn")

        # Store in session state
        if project_name:
            st.session_state.doc_config = {
                "project_name": project_name,
                "format": doc_format,
                "include_summary": include_summary,
                "include_details": include_details,
                "include_failures": include_failures,
                "include_charts": include_charts
            }

    with preview_tab:
        if hasattr(st.session_state, 'test_results') and st.session_state.test_results:
            st.subheader("Results Preview")

            # File selector if multiple files
            if len(st.session_state.test_results) > 1:
                st.session_state.current_file = st.selectbox("Select Result File",
                                                           list(st.session_state.test_results.keys()))

            # Get current file data
            current_data = st.session_state.test_results[st.session_state.current_file]

            # Display summary
            st.info(f"**File:** {st.session_state.current_file}")

            col1, col2, col3, col4, col5 = st.columns(5)
            # Safely handle cases where summary might be an integer or a dictionary
            summary = current_data.get('summary', {})
            if isinstance(summary, dict):
                # Normal case - summary is a dictionary with keys
                with col1:
                    st.metric("Total Tests", summary.get('total', 0))
                with col2:
                    st.metric("Passed", summary.get('passed', 0))
                with col3:
                    st.metric("Failed", summary.get('failed', 0))
                with col4:
                    st.metric("Skipped", summary.get('skipped', 0))
                with col5:
                    st.metric("Duration", f"{summary.get('execution_time', 0):.2f}s")
            else:
                # Case where summary is an integer (likely representing total tests)
                with col1:
                    st.metric("Total Tests", summary if isinstance(summary, (int, float)) else 0)
                with col2:
                    st.metric("Passed", current_data.get('passed', 0))
                with col3:
                    st.metric("Failed", current_data.get('failed', 0))
                with col4:
                    st.metric("Skipped", current_data.get('skipped', 0))
                with col5:
                    st.metric("Duration", f"{current_data.get('execution_time', 0):.2f}s")

            # Create visualization if available
            if VISUALIZATION_AVAILABLE:
                with st.expander("Test Result Visualization", expanded=True):
                    viz_buffer = create_visualization(current_data)
                    if viz_buffer:
                        st.image(viz_buffer, use_container_width=True)

            # Show test details
            with st.expander("Test Details", expanded=False):
                if current_data['tests']:
                    test_df = pd.DataFrame([
                        {
                            "Name": t.get('name', 'Unknown'),
                            "Status": t.get('status', 'unknown').upper(),
                            "Duration (s)": t.get('time', 0),
                            "Class/Suite": t.get('classname', '')
                        }
                        for t in current_data['tests']
                    ])

                    # Apply conditional formatting
                    def highlight_status(val):
                        if val == 'PASSED':
                            return 'background-color: #d4edda; color: #155724'
                        elif val == 'FAILED':
                            return 'background-color: #f8d7da; color: #721c24'
                        elif val == 'SKIPPED':
                            return 'background-color: #fff3cd; color: #856404'
                        else:
                            return ''

                    styled_df = test_df.style.map(highlight_status, subset=['Status'])
                    st.dataframe(styled_df)
                else:
                    st.info("No test details available")

            # Show failures if any
            failures = [t for t in current_data['tests'] if t.get('status') == 'failed']
            if failures:
                with st.expander("Failure Details", expanded=False):
                    for i, failure in enumerate(failures):
                        st.markdown(f"### Failure {i+1}: {failure.get('name', 'Unknown')}")

                        if 'message' in failure:
                            st.markdown(f"**Message:** {failure['message']}")

                        if 'stacktrace' in failure:
                            with st.expander("Stack Trace", expanded=False):
                                st.code(failure['stacktrace'])

    # Generate documentation button
    st.markdown("---")
    col1, col2 = st.columns([2, 1])

    with col1:
        generate_button = st.button("Generate Documentation",
                                  disabled=not (hasattr(st.session_state, 'test_results')
                                              and hasattr(st.session_state, 'doc_config')))

    # Generate documentation
    if generate_button and hasattr(st.session_state, 'test_results') and hasattr(st.session_state, 'doc_config'):
        with st.spinner("Generating documentation..."):
            try:
                # Get configuration
                config = st.session_state.doc_config
                project_name = config['project_name']
                doc_format = config['format']

                # Combine all test data if multiple files
                if len(st.session_state.test_results) > 1:
                    combined_data = {
                        "tests": [],
                        "summary": {
                            "total": 0,
                            "passed": 0,
                            "failed": 0,
                            "skipped": 0,
                            "execution_time": 0
                        }
                    }

                    for file_data in st.session_state.test_results.values():
                        # Add tests
                        combined_data["tests"].extend(file_data["tests"])

                        # Update summary
                        for key in combined_data["summary"]:
                            combined_data["summary"][key] += file_data["summary"].get(key, 0)

                    data = combined_data
                else:
                    # Use current data
                    data = st.session_state.test_results[st.session_state.current_file]

                # Generate documentation based on format
                if doc_format == "HTML":
                    doc_content = generate_html_report(data, project_name)
                    mime_type = OUTPUT_FORMATS[doc_format]["mime"]
                elif doc_format == "Markdown":
                    doc_content = generate_markdown_report(data, project_name)
                    mime_type = OUTPUT_FORMATS[doc_format]["mime"]
                elif doc_format == "PDF":
                    doc_content = generate_pdf_report(data, project_name)
                    mime_type = OUTPUT_FORMATS[doc_format]["mime"]
                    if not doc_content:
                        st.error("Failed to generate PDF. Please install reportlab or choose another format.")
                        doc_format = "Markdown"  # Fallback
                        doc_content = generate_markdown_report(data, project_name)
                        mime_type = OUTPUT_FORMATS["Markdown"]["mime"]
                elif doc_format == "Word":
                    doc_content = generate_docx_report(data, project_name)
                    mime_type = OUTPUT_FORMATS[doc_format]["mime"]
                    if not doc_content:
                        st.error("Failed to generate Word document. Please install python-docx or choose another format.")
                        doc_format = "Markdown"  # Fallback
                        doc_content = generate_markdown_report(data, project_name)
                        mime_type = OUTPUT_FORMATS["Markdown"]["mime"]

                # File extension
                extension = OUTPUT_FORMATS[doc_format]["extension"]

                # Create download name with timestamp
                timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
                filename = f"{project_name.lower().replace(' ', '_')}_{timestamp}.{extension}"

                # Display success message and download button
                st.success(f"Documentation generated successfully!")

                # Send success notification
                if NOTIFICATIONS_AVAILABLE:
                    test_count = data['summary']['total']
                    pass_rate = (data['summary']['passed'] / test_count * 100) if test_count > 0 else 0
                    notifications.add_notification(
                        module_name="auto_documentation",
                        status="success",
                        message=f"Documentation successfully generated in {doc_format} format",
                        details=f"Generated documentation for {test_count} tests with {pass_rate:.1f}% pass rate. Format: {doc_format}.",
                        action_steps=["Download the documentation", "Share with your team"]
                    )

                # Binary formats need special handling
                if doc_format in ["PDF", "Word"]:
                    st.download_button(
                        label=f"Download {doc_format} Documentation",
                        data=doc_content,
                        file_name=filename,
                        mime=mime_type
                    )
                else:
                    st.download_button(
                        label=f"Download {doc_format} Documentation",
                        data=doc_content,
                        file_name=filename,
                        mime=mime_type
                    )

                # Update execution metrics in main UI's session state
                if 'execution_metrics' in st.session_state:
                    st.session_state.execution_metrics['tests_executed'] += data['summary']['total']
                    st.session_state.execution_metrics['successful_tests'] += data['summary']['passed']

                    st.session_state.execution_metrics['failed_tests'] += data['summary']['failed']
                    st.session_state.execution_metrics['execution_time'] += data['summary']['execution_time']

            except Exception as e:
                error_message = f"Failed to generate documentation: {str(e)}"
                st.error(error_message)

                # Send failure notification
                if NOTIFICATIONS_AVAILABLE:
                    notifications.add_notification(
                        module_name="auto_documentation",
                        status="error",
                        message="Documentation generation failed",
                        details=error_message,
                        action_steps=[
                            "Check if the test result files are valid",
                            "Verify you have the required dependencies installed",
                            "Try a different output format"
                        ]
                    )

if __name__ == "__main__":
    # This allows running the module directly for testing
    show_ui()

