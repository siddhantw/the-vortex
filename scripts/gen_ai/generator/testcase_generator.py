import re
import os
import sys

# Add the parent directory to sys.path for imports
current_dir = os.path.dirname(os.path.abspath(__file__))
parent_dir = os.path.dirname(current_dir)
grandparent_dir = os.path.dirname(parent_dir)
if parent_dir not in sys.path:
    sys.path.insert(0, parent_dir)
if grandparent_dir not in sys.path:
    sys.path.insert(0, grandparent_dir)

try:
    from azure_openai_client import AzureOpenAIClient
except ImportError:
    # Fallback import path
    sys.path.insert(0, os.path.join(grandparent_dir, 'gen_ai'))
    from azure_openai_client import AzureOpenAIClient

import json
import datetime
import concurrent.futures
import logging
import time
from docx import Document
from tqdm import tqdm
from typing import List, Dict, Any

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[logging.FileHandler("testcase_generator.log"), logging.StreamHandler()]
)
logger = logging.getLogger("TestCaseGenerator")

DEFAULT_OPENAI_API_KEY = ("")

# Enhanced function to extract test scenarios and other test-related data from requirements
def extract_test_components(requirements_list: List[Dict[str, Any]]) -> Dict[str, List[str]]:
    """
    Enhanced extraction of test scenarios and components from raw requirements with better filtering
    """
    test_data = {
        "test_scenarios": [],
        "test_steps": [],
        "functional_flows": [],
        "acceptance_criteria": [],
        "ambiguous": [],
        "brand_specific_requirements": {},
        "filtered_irrelevant": [],  # Track what was filtered out
        "business_entities": [],    # Track business-relevant entities
        "test_categories": {        # Categorize test types
            "functional": [],
            "non_functional": [],
            "integration": [],
            "security": [],
            "performance": [],
            "usability": []
        }
    }

    # Enhanced brand detection with more comprehensive mapping
    brands = {
        "bluehost": {
            "identifiers": ["bluehost", "bhcom", "bh.com", "blue host", "bluehost.com"],
            "features": ["wordpress hosting", "shared hosting", "vps hosting", "dedicated hosting",
                        "wordpress", "website builder", "domain registration", "woocommerce",
                        "ecommerce", "cpanel", "ssl certificates", "web hosting"],
            "test_priorities": ["hosting_plans", "wordpress_integration", "cpanel_functionality"]
        },
        "domain.com": {
            "identifiers": ["domain.com", "dcom", "domain"],
            "features": ["domain registration", "domain transfer", "domain privacy",
                        "web hosting", "email hosting", "ssl certificates", "website builder"],
            "test_priorities": ["domain_search", "domain_registration", "domain_management"]
        },
        "hostgator": {
            "identifiers": ["hostgator", "hg", "hgcom", "host gator", "hostgator.com"],
            "features": ["shared hosting", "wordpress hosting", "vps hosting", "dedicated hosting",
                        "website builder", "domain registration", "reseller hosting", "cpanel", "ssl"],
            "test_priorities": ["hosting_plans", "reseller_features", "migration_tools"]
        },
        "network solutions": {
            "identifiers": ["network solutions", "netsol", "nsol", "ncom", "networksolutions.com"],
            "features": ["domain registration", "web hosting", "email hosting", "ssl certificates",
                        "dns management", "website builder", "ecommerce", "online marketing"],
            "test_priorities": ["domain_services", "dns_management", "enterprise_features"]
        },
        "register.com": {
            "identifiers": ["register.com", "rcom", "register"],
            "features": ["domain registration", "web hosting", "email hosting", "website builder",
                        "ecommerce", "ssl certificates", "online marketing"],
            "test_priorities": ["domain_registration", "marketing_tools", "ecommerce_integration"]
        },
        "web.com": {
            "identifiers": ["web.com", "wcom"],
            "features": ["website builder", "ecommerce", "digital marketing", "lead generation",
                        "domain registration", "web hosting", "ssl certificates", "seo"],
            "test_priorities": ["website_builder", "marketing_automation", "lead_management"]
        }
    }

    # Enhanced irrelevant content patterns
    irrelevant_patterns = [
        r'\b[A-Z][a-z]+ [A-Z][a-z]+\b(?!\s+(?:hosting|domain|ssl|certificate|website))',  # Names not followed by business terms
        r'\bwritten by\s+[A-Z][a-z]+',
        r'\bversion\s+\d+\.\d+',
        r'\bpage\s+\d+\s+of\s+\d+',
        r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',  # Email addresses
        r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b',  # Phone numbers
        r'\bcopyright\s+©?\s*\d{4}',
        r'\ball rights reserved',
        r'\bconfidential\b',
        r'\b\[.*?\]\b',  # Bracketed placeholders
        r'\b<.*?>\b',    # Angle bracket placeholders
        r'\b\{.*?\}\b'   # Curly bracket placeholders
    ]

    # Business-relevant keywords for better filtering
    business_keywords = [
        'domain', 'hosting', 'website', 'ssl', 'certificate', 'dns', 'server',
        'email', 'backup', 'security', 'performance', 'uptime', 'bandwidth',
        'storage', 'database', 'wordpress', 'ecommerce', 'cpanel', 'ftp',
        'migration', 'transfer', 'registration', 'renewal', 'billing', 'payment',
        'support', 'ticket', 'dashboard', 'control panel', 'user', 'customer',
        'account', 'login', 'authentication', 'authorization', 'api', 'integration'
    ]

    # Enhanced test scenario patterns
    test_scenario_patterns = [
        r'\bverify\s+that\b.*',
        r'\bensure\s+that\b.*',
        r'\btest\s+(?:that|whether|if)\b.*',
        r'\bcheck\s+(?:that|whether|if)\b.*',
        r'\bvalidate\s+that\b.*',
        r'\bconfirm\s+that\b.*'
    ]

    def is_business_relevant(text: str) -> bool:
        """Check if text contains business-relevant content"""
        text_lower = text.lower()
        business_score = sum(1 for keyword in business_keywords if keyword in text_lower)
        functional_indicators = ['shall', 'must', 'will', 'should', 'can', 'user', 'system']
        functional_score = sum(1 for indicator in functional_indicators if indicator in text_lower)
        return business_score >= 2 or functional_score >= 2

    def filter_irrelevant_content(text: str) -> str:
        """Filter out irrelevant content from text"""
        filtered_text = text
        for pattern in irrelevant_patterns:
            filtered_text = re.sub(pattern, '', filtered_text, flags=re.IGNORECASE)
        return re.sub(r'\s+', ' ', filtered_text).strip()

    def categorize_test_type(text: str) -> str:
        """Categorize the type of test based on content"""
        text_lower = text.lower()

        if any(word in text_lower for word in ['performance', 'speed', 'load', 'response time', 'throughput']):
            return 'performance'
        elif any(word in text_lower for word in ['security', 'authentication', 'authorization', 'encryption', 'ssl']):
            return 'security'
        elif any(word in text_lower for word in ['integration', 'api', 'third party', 'external']):
            return 'integration'
        elif any(word in text_lower for word in ['usability', 'user experience', 'interface', 'navigation']):
            return 'usability'
        elif any(word in text_lower for word in ['availability', 'reliability', 'uptime', 'scalability']):
            return 'non_functional'
        else:
            return 'functional'

    # Process each requirement with enhanced filtering and analysis
    for req in requirements_list:
        req_text = req.get("text", "")

        # Skip very short requirements
        if len(req_text) < 10:
            continue

        # Filter irrelevant content first
        original_text = req_text
        filtered_text = filter_irrelevant_content(req_text)

        # Track what was filtered out
        if len(filtered_text) < len(original_text) * 0.8:  # More than 20% was filtered
            test_data["filtered_irrelevant"].append({
                "original": original_text,
                "filtered": filtered_text,
                "reason": "Excessive irrelevant content"
            })

        # Skip if too much content was filtered or not business relevant
        if len(filtered_text) < 5 or not is_business_relevant(filtered_text):
            test_data["filtered_irrelevant"].append({
                "original": original_text,
                "filtered": filtered_text,
                "reason": "Not business relevant"
            })
            continue

        # Use filtered text for further processing
        req_text = filtered_text

        # Enhanced brand identification
        identified_brands = []
        for brand, info in brands.items():
            identifiers = info["identifiers"]
            if any(identifier in req_text.lower() for identifier in identifiers):
                identified_brands.append(brand)

                # Add to brand-specific requirements
                if brand not in test_data["brand_specific_requirements"]:
                    test_data["brand_specific_requirements"][brand] = []
                test_data["brand_specific_requirements"][brand].append(req_text)

                # Generate brand-specific test scenarios based on features
                for feature in info["features"]:
                    if feature in req_text.lower():
                        feature_scenario = f"Verify {feature} functionality for {brand}"
                        if feature_scenario not in test_data["test_scenarios"]:
                            test_data["test_scenarios"].append(feature_scenario)

        # If no specific brand mentioned, add to generic requirements
        if not identified_brands:
            if "all brands" not in test_data["brand_specific_requirements"]:
                test_data["brand_specific_requirements"]["all brands"] = []
            test_data["brand_specific_requirements"]["all brands"].append(req_text)

        # Enhanced scenario extraction using patterns
        is_test_scenario = False
        for pattern in test_scenario_patterns:
            if re.search(pattern, req_text, re.IGNORECASE):
                test_data["test_scenarios"].append(req_text)
                is_test_scenario = True
                break

        # Enhanced ambiguity checking
        ambiguity_score = req.get("ambiguity", 0.0)
        completeness_score = req.get("completeness", 0.0)

        # Only add high-quality requirements to scenarios
        if not is_test_scenario and ambiguity_score < 0.3 and completeness_score > 0.6:
            test_data["test_scenarios"].append(req_text)

            # Extract potential test steps
            sentences = re.split(r'[.!?]', req_text)
            for sentence in sentences:
                clean_sentence = sentence.strip()
                if len(clean_sentence) > 10 and is_business_relevant(clean_sentence):
                    test_data["test_steps"].append(clean_sentence)

        # Categorize test type
        test_category = categorize_test_type(req_text)
        test_data["test_categories"][test_category].append(req_text)

        # Enhanced flow and criteria detection
        if any(word in req_text.lower() for word in ["flow", "process", "workflow", "sequence"]):
            test_data["functional_flows"].append(req_text)

        if any(word in req_text.lower() for word in ["should", "must", "will", "shall", "criteria"]):
            test_data["acceptance_criteria"].append(req_text)

        # Enhanced ambiguous requirement handling
        if ambiguity_score > 0.5:
            test_data["ambiguous"].append({
                "text": req_text,
                "ambiguity_score": ambiguity_score,
                "issues": req.get("analysis_results", {}).get("ambiguity", {})
            })

    # Generate intelligent test scenarios if none found
    if not test_data["test_scenarios"]:
        # Create brand-specific scenarios based on identified brands and their priorities
        for brand, requirements in test_data["brand_specific_requirements"].items():
            if brand != "all brands" and brand in brands:
                brand_info = brands[brand]
                for priority in brand_info.get("test_priorities", []):
                    scenario = f"Verify {priority.replace('_', ' ')} for {brand}"
                    test_data["test_scenarios"].append(scenario)

        # Add generic high-priority scenarios if still empty
        if not test_data["test_scenarios"]:
            generic_scenarios = [
                "Verify user account registration and login functionality",
                "Verify domain search and registration process",
                "Verify hosting plan selection and purchase flow",
                "Verify SSL certificate installation and activation",
                "Verify billing and payment processing",
                "Verify customer support ticket creation and management",
                "Verify control panel access and functionality",
                "Verify website builder tools and features",
                "Verify email account setup and management",
                "Verify backup and restore functionality"
            ]
            test_data["test_scenarios"].extend(generic_scenarios)

    return test_data

class TestCaseGenerator:
    def __init__(self, analysis: dict, config: dict):
        self.analysis = analysis
        self.config = config
        self.use_llm = config.get('use_llm', False) if config else False
        self.max_parallel_tasks = config.get('max_parallel_tasks', 3) if config else 3
        self.use_memory_optimization = config.get('use_memory_optimization', True) if config else True

        # Auto-select best Azure OpenAI model based on complexity
        if config and config.get('auto_select_model', True):
            complexity = self._estimate_requirement_complexity(analysis)
            if complexity > 0.7:  # High complexity
                self.llm_model = 'gpt-4o'  # Use GPT-4o for complex requirements
            elif complexity > 0.4:  # Medium complexity
                self.llm_model = 'gpt-4'   # Use GPT-4 for medium complexity
            else:  # Low complexity
                self.llm_model = 'gpt-35-turbo'  # Use GPT-3.5 Turbo for simple requirements
            logger.info(f"Auto-selected Azure OpenAI model {self.llm_model} based on complexity score: {complexity:.2f}")
        else:
            # Default to GPT-4 if no auto-selection or fallback from config
            self.llm_model = config.get('llm_model', 'gpt-4') if config else 'gpt-4'

        # Configure Azure OpenAI client with enhanced error handling and fallback
        self.azure_openai_config = {
            'azure_endpoint': os.environ.get('AZURE_OPENAI_ENDPOINT'),
            'api_key': os.environ.get('AZURE_OPENAI_API_KEY', "5e98b3558f5d4dcebe68f8ca8a3352b7"),
            'api_version': config.get('azure_api_version', '2024-10-21') if config else '2024-10-21',
            'deployment_name': self.llm_model  # Use the selected model as deployment name
        }

        # Initialize Azure OpenAI client for test case generation
        try:
            self.azure_client = AzureOpenAIClient(
                azure_endpoint=self.azure_openai_config['azure_endpoint'],
                api_key=self.azure_openai_config['api_key'],
                api_version=self.azure_openai_config['api_version'],
                deployment_name=self.azure_openai_config['deployment_name']
            )
            self.has_azure_openai = True
            logger.info(f"Azure OpenAI client initialized successfully with model: {self.llm_model}")
        except Exception as e:
            logger.warning(f"Failed to initialize Azure OpenAI client: {e}. AI-enhanced generation will be disabled.")
            self.azure_client = None
            self.has_azure_openai = False

        self.output_dir = config.get('output_dir', 'generated_tests') if config else 'generated_tests'
        os.makedirs(self.output_dir, exist_ok=True)
        self.test_cases = []
        self.test_case_count = 0
        self.test_case_name = None
        self.test_case_description = None
        self.test_case_preconditions = None
        self.test_case_steps = None
        self.test_case_expected = None
        self.test_case_postconditions = None
        self.test_case_priority = None
        self.test_case_severity = None
        self.test_case_automation_feasibility = None
        self.test_case_automatable = None
        self.test_case_type = None
        self.test_case_name_robot = None
        self.test_case_conditions_steps = None
        self.test_case_test_data = None
        self.test_case_brand = None

        # Cache for expensive operations
        self._cache = {}

    def _estimate_requirement_complexity(self, analysis):
        """Estimate complexity of requirements to choose appropriate AI model"""
        complexity_score = 0.0

        # Count number of identified items
        num_scenarios = len(analysis.get("test_scenarios", []))
        num_flows = len(analysis.get("functional_flows", []))
        num_criteria = len(analysis.get("acceptance_criteria", []))

        # Text length and complexity factors
        all_text = ""
        for key, values in analysis.items():
            if isinstance(values, list):
                all_text += " ".join(str(v) for v in values)
            elif isinstance(values, str):
                all_text += values

        text_length = len(all_text)
        num_entities = len(set(re.findall(r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\b', all_text)))

        # Calculate complexity score (normalized to 0-1 range)
        if text_length > 0:
            complexity_score = min(1.0, (
                0.3 * min(1.0, num_scenarios / 20) +
                0.2 * min(1.0, num_flows / 15) +
                0.2 * min(1.0, num_criteria / 15) +
                0.1 * min(1.0, text_length / 10000) +
                0.2 * min(1.0, num_entities / 50)
            ))

        return complexity_score

    def _make_test_name(self, component, func, case_type, idx):
        brand = self._extract_brands(func)
        brand_part = brand.replace(' ', '').replace(',', '_').replace('and', '_').upper() if brand != 'All Brands' else 'ALL'
        comp = (component or 'GEN').upper().replace(' ', '_')
        func_part = re.sub(r'[^A-Za-z0-9]+', '_', func.strip().split()[0].upper())[:20]
        case_type = case_type.upper()
        return f"TC_{brand_part}_{comp}_{case_type}_{str(idx).zfill(3)}"

    def _llm_generate(self) -> list:
        # Initialize Azure OpenAI client
        try:
            azure_client = AzureOpenAIClient()
        except Exception as e:
            raise ValueError(f"Failed to initialize Azure OpenAI client: {e}")

        import uuid
        mcp_context = self._make_mcp_context()
        prompt = (
            "You are an expert QA engineer. Given the following Model Context Protocol (MCP) for a software system, "
            "generate a comprehensive suite of Robot Framework test cases. "
            "For each test, include: name, preconditions, steps, expected results, postconditions, priority, and severity. "
            "Include positive, negative, and edge cases for each functional flow and acceptance criterion. "
            "Use the naming convention TC_<COMPONENT>_<FUNCTION>_<POS|NEG|EDGE>_<ID>. "
            "Output only valid Robot Framework test case syntax.\n\n"
            f"MCP Context:\n{mcp_context}\n"
        )
        prompt += "Test Cases:\n"

        response = azure_client.chat_completion_create(
            model=self.llm_model,
            messages=[{"role": "user", "content": prompt}],
            max_tokens=3000,
            temperature=0.2,
            top_p=1,
            frequency_penalty=0,
            presence_penalty=0,
            stop=["\n\n"]
        )

        content = response["choices"][0]["message"]["content"]
        return [{"llm_output": content}]

    def _make_mcp_context(self):
        """Create a context string from the analysis data for the language model"""
        context = []

        # Add system description if available
        if "system_description" in self.analysis:
            context.append(f"System Description: {self.analysis['system_description']}")

        # Add scenarios
        if "test_scenarios" in self.analysis and self.analysis["test_scenarios"]:
            context.append("Test Scenarios:")
            for i, scenario in enumerate(self.analysis["test_scenarios"]):
                context.append(f"  {i+1}. {scenario}")

        # Add functional flows
        if "functional_flows" in self.analysis and self.analysis["functional_flows"]:
            context.append("Functional Flows:")
            for i, flow in enumerate(self.analysis["functional_flows"]):
                context.append(f"  {i+1}. {flow}")

        # Add acceptance criteria
        if "acceptance_criteria" in self.analysis and self.analysis["acceptance_criteria"]:
            context.append("Acceptance Criteria:")
            for i, criteria in enumerate(self.analysis["acceptance_criteria"]):
                context.append(f"  {i+1}. {criteria}")

        # Add test steps if available
        if "test_steps" in self.analysis and self.analysis["test_steps"]:
            context.append("Test Steps:")
            for i, step in enumerate(self.analysis["test_steps"]):
                context.append(f"  {i+1}. {step}")

        return "\n".join(context)

    def _extract_brands(self, text: str) -> str:
        brand_map = {
            'bluehost.com': 'BHCOM', 'bluehost': 'BHCOM', 'bhcom': 'BHCOM',
            'domain.com': 'DCOM', 'dcom': 'DCOM',
            'hostgator.com': 'HGCOM', 'hostgator': 'HG', 'hg': 'HG',
            'network solutions': 'NCOM', 'nsol': 'NCOM', 'ncom': 'NCOM',
            'register.com': 'RCOM', 'rcom': 'RCOM',
            'web.com': 'WCOM', 'wcom': 'WCOM',
        }
        found = set()
        text_lower = text.lower()
        for key, code in brand_map.items():
            if key in text_lower:
                found.add(code)

        # If no brands found, or if ALL major brands are found, return "All Brands"
        if not found:
            return 'All Brands'

        # Define all major brand codes
        all_major_brands = {'BHCOM', 'DCOM', 'HG', 'NCOM', 'RCOM', 'WCOM'}

        # If all major brands are found or 5+ brands, return "All Brands" instead of listing them all
        if found == all_major_brands or len(found) >= 5:
            return 'All Brands'

        return ', '.join(sorted(found))

    def _get_brand_code(self, brand_name: str) -> str:
        """Convert a brand name to its code representation."""
        brand_map = {
            'bluehost': 'BHCOM',
            'bluehost.com': 'BHCOM',
            'bhcom': 'BHCOM',
            'domain.com': 'DCOM',
            'dcom': 'DCOM',
            'hostgator': 'HG',
            'hostgator.com': 'HGCOM',
            'hg': 'HG',
            'network solutions': 'NCOM',
            'nsol': 'NCOM',
            'netsol': 'NCOM',
            'ncom': 'NCOM',
            'register.com': 'RCOM',
            'register': 'RCOM',
            'rcom': 'RCOM',
            'web.com': 'WCOM',
            'wcom': 'WCOM',
        }

        # Try to find direct match (case insensitive)
        brand_name_lower = brand_name.lower()
        if brand_name_lower in brand_map:
            return brand_map[brand_name_lower]

        # Try to find partial match
        for key, code in brand_map.items():
            if key in brand_name_lower or brand_name_lower in key:
                return code

        # Default to "All Brands" if no match found
        return "All Brands"

    def _extract_functionalities(self, analysis):
        items = set()
        for k in [
            "functionalities", "functional_flows", "acceptance_criteria", "integration_points", "constraints",
            "non_functional_requirements", "system_description"
        ]:
            for v in analysis.get(k, []):
                if v and len(v.strip()) > 8:
                    for part in re.split(r'[.;]|\band\b|\bthen\b|\bbut\b|\bor\b|\balso\b|as well as|followed by|after|before|when|if|unless|until|once|provided|where|while|during|upon|with|without|in case|in order to|so that|such that|for|to|by|from|on|at|of|in|into|onto|over|under|above|below|through|across|between|among|against|along|around|about|like|as|via|per|according to|including|excluding|except|besides|plus|minus|versus|vs|etc\.', v, flags=re.I):
                        part = part.strip()
                        if len(part) > 8:
                            for subpart in part.split(','):
                                subpart = subpart.strip()
                                if len(subpart) > 8:
                                    items.add(subpart)
        expanded = set()
        for item in items:
            if re.match(r"^\d+\. ", item):
                for step in re.split(r'\d+\. ', item):
                    step = step.strip()
                    if len(step) > 8:
                        expanded.add(step)
            else:
                expanded.add(item)
        final = set()
        for item in expanded:
            for match in re.finditer(r'\b([A-Za-z]+)\b\s+(the|a|an)?\s*([A-Za-z0-9_\- ]+)', item):
                verb, _, obj = match.groups()
                if verb and obj and len(obj.strip()) > 2:
                    final.add(f"{verb.strip().capitalize()} {obj.strip()}")
            final.add(item)
        verbs = set()
        nouns = set()
        for item in final:
            for match in re.finditer(r'\b([A-Za-z]+)\b\s+(the|a|an)?\s*([A-Za-z0-9_\- ]+)', item):
                verb, _, obj = match.groups()
                if verb and obj:
                    verbs.add(verb.strip().capitalize())
                    nouns.add(obj.strip())
        for v in verbs:
            for n in nouns:
                if len(n) > 2:
                    final.add(f"{v} {n}")
        return sorted(final)

    def _generate_realistic_test_data(self, scenario, steps, type="POS"):
        data_entries = []
        all_text = scenario + " " + " ".join(steps if isinstance(steps, list) else [steps])
        fields = re.findall(r'\b(username|email|password|name|address|phone|zip|postal code|credit card|payment|amount|date|id|code|token|account|domain|hostname|url|path|file|directory|value|number|quantity|price|rate|description|title)\b', all_text, re.I)
        unique_fields = set(f.lower() for f in fields)
        data_templates = {
            "username": ["user123", "admin_user", "test_user"] if type == "POS" else ["", "a", "user with spaces", "user'with\"quotes", "1234567890"*10],
            "email": ["test@example.com", "user@company.org"] if type == "POS" else ["invalid-email", "test@", "@domain.com", "test@.com", ""],
            "password": ["SecurePass123!"] if type == "POS" else ["", "short", "password", "12345678"],
            "name": ["John Smith", "Jane Doe"] if type == "POS" else ["", "J", "Very long name that exceeds the maximum allowed character limit for names"],
            "address": ["123 Main St, City, State 12345"] if type == "POS" else ["", "Invalid address format"],
            "phone": ["555-123-4567", "+1 (555) 987-6543"] if type == "POS" else ["555", "not-a-number", ""],
            "credit card": ["4111111111111111", "5555555555554444"] if type == "POS" else ["4111", "invalid-card", ""],
            "payment": ["Credit Card", "PayPal", "Bank Transfer"] if type == "POS" else ["Invalid Payment Type", ""],
            "amount": ["100.00", "1250.75"] if type == "POS" else ["0", "-100", "100,000,000,000.00", "0.001"],
            "date": ["2025-05-20", "05/20/2025"] if type == "POS" else ["invalid-date", "0000-00-00", ""],
            "domain": ["example.com", "test-domain.org"] if type == "POS" else ["invalid domain", ".com", "example..com"],
            "quantity": ["1", "10", "100"] if type == "POS" else ["0", "-1", "1000000", ""],
            "url": ["https://example.com", "http://test.org/path"] if type == "POS" else ["http//missing-colon", "ftp:/invalid", ""]
        }
        if "bluehost" in all_text.lower():
            data_templates["domain"] = ["example-bluehost.com", "mybusiness-bh.org"] if type == "POS" else ["invalid_bh$.com", ".bluehost"]
        elif "hostgator" in all_text.lower():
            data_templates["domain"] = ["example-hg.com", "mybusiness-hostgator.org"] if type == "POS" else ["invalid_hg$.com", ".hostgator"]
        elif "domain.com" in all_text.lower():
            data_templates["domain"] = ["premium-domain.com", "business-domain.org"] if type == "POS" else ["invalid_dcom$.com", ".domain"]
        for field in unique_fields:
            template_key = next((k for k in data_templates.keys() if k in field), None)
            if template_key:
                values = data_templates[template_key]
                if type == "POS":
                    values = values[:2]
                else:
                    values = values[-2:]
                for val in values:
                    data_entries.append(f"{field}={val}")
        if type == "EDGE":
            numeric_fields = {"quantity", "amount", "price", "number"}
            found_numeric = any(nf in unique_fields for nf in numeric_fields)
            if found_numeric:
                data_entries.append("Boundary values: 0, 1, -1, max allowed value, min allowed value")
            length_sensitive = {"name", "username", "password", "address", "description", "title"}
            found_length = any(ls in unique_fields for ls in length_sensitive)
            if found_length:
                data_entries.append("Length tests: minimum allowed length, maximum allowed length, maximum+1")
        if "login" in all_text.lower() or "authenticate" in all_text.lower():
            if type == "POS":
                data_entries.append("Valid credentials (username/password combination)")
            elif type == "NEG":
                data_entries.append("Invalid credentials, locked account, expired account")
            else:
                data_entries.append("Account with minimum permissions, account near expiration")
        if "purchase" in all_text.lower() or "payment" in all_text.lower() or "order" in all_text.lower():
            if type == "POS":
                data_entries.append("Valid product SKUs/IDs, valid payment details")
            elif type == "NEG":
                data_entries.append("Invalid product ID, insufficient funds, payment declined")
            else:
                data_entries.append("Highest priced item, lowest priced item, bulk order maximums")
        if not data_entries:
            if type == "POS":
                data_entries.append("Valid input data per business requirements")
            elif type == "NEG":
                data_entries.append("Invalid input data that violates business rules")
            else:
                data_entries.append("Boundary values, minimum/maximum allowed values")
        return "\n".join(data_entries)

    def _determine_automation_feasibility(self, scenario, steps, priority="Medium", severity="Medium"):
        all_text = scenario + " " + " ".join(steps if isinstance(steps, list) else [steps])
        manual_indicators = [
            "visual", "appearance", "look and feel", "ux", "user experience",
            "performance", "speed", "responsive", "usability testing",
            "verify email received", "check email content", "physical", "hardware",
            "manual validation", "human verification", "captcha",
            "subjective", "aesthetic", "compare", "human judgment"
        ]
        automation_indicators = [
            "login", "create account", "form submission", "validation message",
            "error message", "field validation", "navigation", "search functionality",
            "api request", "database check", "response", "status code", "json",
            "element", "button", "click", "input", "submit", "link", "menu"
        ]
        manual_score = sum(2 if mi in all_text.lower() else 0 for mi in manual_indicators)
        automation_score = sum(1 if ai in all_text.lower() else 0 for ai in automation_indicators)
        if isinstance(steps, list) and len(steps) > 10:
            automation_score -= 3
        ambiguous_terms = ["may", "might", "possibly", "perhaps", "if applicable", "sometimes", "check if"]
        ambiguity_score = sum(3 if at in all_text.lower() else 0 for at in ambiguous_terms)

        # Definitive cases based on indicators
        if manual_score > automation_score or ambiguity_score >= 3:
            return "No"
        elif automation_score > 0:
            return "Yes"

        # For cases that would have returned "Maybe", make a decision based on priority and severity
        # Convert priority and severity to uppercase for case-insensitive comparison
        priority = priority.upper() if isinstance(priority, str) else "MEDIUM"
        severity = severity.upper() if isinstance(severity, str) else "MEDIUM"

        # For High/Critical priority or severity, prefer automation
        if priority in ["HIGH", "CRITICAL"] or severity in ["HIGH", "CRITICAL", "MAJOR"]:
            return "Yes"  # Prioritize automation for high-value test cases
        else:
            return "No"   # Default to manual testing for lower priority/severity cases

    def _craft_expected_result(self, scenario, steps, case_type, acceptance_criteria=None):
        if acceptance_criteria:
            matched_criteria = []
            for criteria in acceptance_criteria:
                if criteria in scenario or any(criteria in step for step in steps if isinstance(steps, list) and step):
                    return f"Expected: {criteria}"
        all_text = scenario + " " + " ".join(steps if isinstance(steps, list) else [steps])

        # Try to find explicit expectations in the requirement text
        expects = re.findall(r'(should|must|will|shall|expected to|needs to) ([^.;:]+)', all_text, re.I)
        if expects:
            return f"Expected: System {expects[0][0]} {expects[0][1]}"
        if case_type == "POS":
            if "login" in all_text.lower():
                return "Expected:\n- User is successfully authenticated\n- User is redirected to the dashboard/homepage\n- User session is correctly established\n- Appropriate welcome message is displayed"
            elif "create" in all_text.lower() or "add" in all_text.lower() or "register" in all_text.lower():
                return "Expected:\n- New entity is successfully created\n- Data is correctly stored in the system\n- Confirmation message is displayed\n- New record appears in relevant listings"
            elif "update" in all_text.lower() or "edit" in all_text.lower() or "modify" in all_text.lower():
                return "Expected:\n- Changes are successfully saved\n- Updated data is reflected in the system\n- Confirmation message is displayed\n- Changed values are visible in the UI"
            elif "delete" in all_text.lower() or "remove" in all_text.lower():
                return "Expected:\n- Entity is successfully removed\n- Data is no longer accessible in the system\n- Confirmation message is displayed\n- Record no longer appears in relevant listings"
            elif "search" in all_text.lower() or "find" in all_text.lower():
                return "Expected:\n- Accurate search results are displayed\n- Results match the search criteria\n- Search metadata (count, time) is accurate\n- Sorting and filtering options work correctly"
            elif "order" in all_text.lower() or "purchase" in all_text.lower() or "checkout" in all_text.lower():
                return "Expected:\n- Order is successfully processed\n- Confirmation number is generated\n- Receipt is displayed/sent\n- Inventory is updated accordingly\n- Payment is processed correctly"
            else:
                return "Expected: Operation completes successfully according to business requirements with appropriate confirmation"
        elif case_type == "NEG":
            if "login" in all_text.lower():
                return "Expected:\n- System displays specific error message\n- Access is denied\n- Security measures are activated if applicable\n- Failed login attempt is logged"
            elif any(term in all_text.lower() for term in ["create", "add", "register", "update", "edit", "modify", "delete", "remove"]):
                return "Expected:\n- Operation fails with appropriate error message\n- Validation failure is clearly explained\n- No data corruption occurs\n- User can correct the issue and retry"
            elif "search" in all_text.lower() or "find" in all_text.lower():
                return "Expected:\n- Appropriate message indicating no results found\n- Invalid search criteria is highlighted\n- Helpful suggestions are provided\n- User can modify search and try again"
            elif "order" in all_text.lower() or "purchase" in all_text.lower() or "checkout" in all_text.lower():
                return "Expected:\n- Order process halts with clear error message\n- No charges are processed\n- Customer data is preserved where appropriate\n- Error details are logged for troubleshooting"
            else:
                return "Expected: System handles invalid input gracefully with appropriate error message; no system crash or data corruption"
        else:
            if "timeout" in all_text.lower() or "performance" in all_text.lower():
                return "Expected: System handles the condition gracefully within performance SLAs; appropriate timeout message if applicable"
            elif "concurrent" in all_text.lower() or "simultaneous" in all_text.lower():
                return "Expected: System maintains data integrity and proper state handling during concurrent operations"
            elif "limit" in all_text.lower() or "maximum" in all_text.lower() or "minimum" in all_text.lower():
                return "Expected: System correctly enforces boundary conditions with appropriate messaging"
            else:
                return "Expected: System handles edge condition appropriately maintaining data integrity and user experience"

    def _generate_scenario_steps(self, scenario, steps_list, requirements, test_type="POS"):
        if steps_list and isinstance(steps_list, list) and len(steps_list) > 0:
            cleaned_steps = []
            for i, step in enumerate(steps_list):
                if not step or len(step.strip()) < 5:
                    continue
                step = step.strip()
                if not re.match(r'^[A-Z][a-z]+', step):
                    common_verbs = ["Navigate", "Click", "Enter", "Verify", "Check", "Select", "Submit", "Confirm"]
                    if "login" in step.lower():
                        step = f"Login with {test_type.lower()} credentials"
                    elif "search" in step.lower():
                        step = f"Search for {test_type.lower()} test data"
                    else:
                        step = f"{common_verbs[i % len(common_verbs)]} {step}"
                cleaned_steps.append(step)
            if cleaned_steps:
                return cleaned_steps
        scenario_lower = scenario.lower()
        generated_steps = []
        if "login" in scenario_lower or "authenticate" in scenario_lower or "sign in" in scenario_lower:
            steps = [
                "Navigate to the login page",
                f"Enter {'valid' if test_type=='POS' else 'invalid'} username",
                f"Enter {'valid' if test_type=='POS' else 'invalid'} password",
                "Click the Login button",
                f"Verify {'successful login' if test_type=='POS' else 'appropriate error message'}"
            ]
            generated_steps.extend(steps)
        elif "register" in scenario_lower or "sign up" in scenario_lower or "create account" in scenario_lower:
            steps = [
                "Navigate to the registration page",
                "Enter user information (name, email, etc.)",
                "Create a password meeting security requirements",
                "Accept terms and conditions",
                "Click the Register/Sign Up button",
                f"Verify {'account creation' if test_type=='POS' else 'validation errors'}"
            ]
            generated_steps.extend(steps)
        elif "checkout" in scenario_lower or "purchase" in scenario_lower or "payment" in scenario_lower:
            steps = [
                "Add products to cart",
                "Navigate to checkout page",
                "Enter shipping information",
                "Select shipping method",
                "Enter payment details",
                "Review order summary",
                "Submit order",
                f"Verify {'order confirmation' if test_type=='POS' else 'payment validation errors'}"
            ]
            generated_steps.extend(steps)
        elif "search" in scenario_lower or "find" in scenario_lower:
            steps = [
                "Navigate to search feature",
                f"Enter {'valid' if test_type=='POS' else 'invalid'} search criteria",
                "Execute search",
                f"Verify {'search results display correctly' if test_type=='POS' else 'no results or error message'}"
            ]
            generated_steps.extend(steps)
        if not generated_steps:
            steps = [
                f"Navigate to the relevant {'page' if 'page' in scenario_lower else 'section'}",
                f"{'Enter required information' if test_type=='POS' else 'Enter invalid data'}",
                "Submit the form/request",
                f"Verify {'successful operation' if test_type=='POS' else 'appropriate error handling'}"
            ]
            generated_steps.extend(steps)
        return generated_steps

    def _generate_branded_test_name(self, scenario, test_type, brand):
        """Generate a descriptive, actionable test case name that clearly identifies the purpose of the test"""
        # Start with identifying the main action/function being tested
        action_verbs = ["verify", "test", "validate", "check", "ensure", "confirm"]
        main_flow = ""

        # Look for key functional actions in the scenario
        action_patterns = [
            r"(log[^\w]+in|authenticate|sign[^\w]+in|access)",
            r"(register|sign[^\w]+up|create[^\w]+account|new[^\w]+user)",
            r"(checkout|purchase|order|buy|payment|transaction)",
            r"(search|find|filter|look[^\w]+for)",
            r"(create|add|insert|upload)",
            r"(update|edit|modify|change)",
            r"(delete|remove|cancel|disable)",
            r"(view|display|show|list|browse)",
            r"(export|download|save)",
            r"(import|upload)",
            r"(validate|verify|check)"
        ]

        for pattern in action_patterns:
            matches = re.findall(pattern, scenario.lower())
            if matches:
                main_action = matches[0]
                # Found the primary action, now extract what is being acted upon
                action_idx = scenario.lower().find(main_action)
                if action_idx >= 0:
                    # Look for a noun or object after the action
                    remaining_text = scenario[action_idx + len(main_action):].strip()
                    # Extract first 3-5 significant words after the action
                    words = [w for w in remaining_text.split()
                             if len(w) > 2 and w.lower() not in ["the", "and", "that", "with", "this", "for", "to"]]
                    object_phrase = " ".join(words[:min(5, len(words))])
                    main_flow = f"{main_action} {object_phrase}"
                    break

        # If we couldn't identify a clear action-object pattern, use the first sentence or phrase
        if not main_flow:
            # Extract first sentence or up to 80 chars
            main_flow = re.split(r'[.!?]', scenario)[0]
            if len(main_flow) > 80:
                words = main_flow.split()
                main_flow = " ".join(words[:min(10, len(words))])

        # Format based on test type
        test_type_prefix = ""
        if test_type == "POS":
            test_type_prefix = "Verify"
        elif test_type == "NEG":
            test_type_prefix = "Verify error handling when"
        else:  # EDGE
            test_type_prefix = "Verify boundary conditions for"

        # Add brand prefix if applicable
        brand_prefix = ""
        if brand and brand != "All Brands":
            brand_prefix = f"[{brand}] "

        # Create final descriptive, actionable name
        result = f"{brand_prefix}{test_type_prefix} {main_flow.strip()}"

        # Capitalize first letter of each word for readability
        result = ' '.join(word.capitalize() if i > 0 else word
                          for i, word in enumerate(result.split()))

        # Ensure reasonable length
        if len(result) > 100:
            result = result[:97] + "..."

        return result.strip()

    def generate(self, scope="Both", test_type=None, components=None) -> list:
        if self.use_llm:
            return self._llm_generate()
        test_cases = []
        analysis = self.analysis
        scenarios = [s for s in analysis.get("test_scenarios", []) if len(s.strip()) > 8 and not s.strip().lower().startswith("ambiguous")]
        steps_list = [s for s in analysis.get("test_steps", []) if len(s.strip()) > 4]
        func_items = [f for f in self._extract_functionalities(analysis) if len(f.strip()) > 8]
        ambiguous = analysis.get("ambiguous", [])
        acceptance_criteria = [a for a in analysis.get("acceptance_criteria", []) if len(a.strip()) > 8]
        flows = [f for f in analysis.get("functional_flows", []) if len(f.strip()) > 8]

        # Process brand-specific requirements
        brand_specific_reqs = analysis.get("brand_specific_requirements", {})
        brand_scenarios = []

        for brand, requirements in brand_specific_reqs.items():
            for req in requirements:
                # Create a brand-specific scenario
                if len(req.strip()) > 8 and not req.strip().lower().startswith("ambiguous"):
                    brand_scenarios.append({"text": req, "brand": brand})

        if not (scenarios or func_items or brand_scenarios):
            print("[DEBUG] No clear scenarios or functionalities extracted. Analysis results:")
            print(f"[DEBUG] Analysis: {analysis}")

        idx = 1
        high_quality_scenarios = []

        # First process regular scenarios
        for scenario in scenarios:
            if test_type and test_type.lower() not in scenario.lower():
                continue
            component = None
            if components:
                for comp in components:
                    if comp.lower() in scenario.lower():
                        component = comp
                        break
            raw_brand = self._extract_brands(scenario)
            brand = raw_brand
            matched_steps = [s for s in steps_list if s in scenario]
            scenario_steps = self._generate_scenario_steps(
                scenario,
                matched_steps if matched_steps else None,
                {"acceptance": acceptance_criteria, "flows": flows}
            )
            automatable = self._determine_automation_feasibility(scenario, scenario_steps)
            for case_type in ["POS", "NEG", "EDGE"]:
                test_data = self._generate_realistic_test_data(scenario, scenario_steps, case_type)
                expected = self._craft_expected_result(scenario, scenario_steps, case_type, acceptance_criteria)
                test_name = self._generate_branded_test_name(scenario, case_type, brand)

                # Create a more descriptive and actionable description
                description = scenario.strip()
                if case_type == "POS":
                    description = f"This test verifies that {scenario.strip().lower()}"
                elif case_type == "NEG":
                    description = f"This test verifies proper error handling when {scenario.strip().lower()}"
                else:  # EDGE
                    description = f"This test validates boundary conditions when {scenario.strip().lower()}"

                tc_id = self._make_test_name(component, scenario, case_type, idx)
                conditions_steps_formatted = '\n'.join(f"{i+1}. {s}" for i, s in enumerate(scenario_steps))
                test_cases.append({
                    "TC id": tc_id,
                    "brand": brand,
                    "name": test_name,
                    "description": description,
                    "conditions/steps": conditions_steps_formatted,
                    "test data": test_data,
                    "priority": "High" if case_type == "POS" else ("Medium" if case_type == "EDGE" else "Low"),
                    "severity": "Critical" if case_type == "POS" else ("Major" if case_type == "EDGE" else "Minor"),
                    "automation feasibility": automatable,
                    "expected result": expected,
                    "_automatable": automatable == "Yes",
                    "name_robot": f"{tc_id} {test_name[:40]}",
                    "steps": scenario_steps,
                    "preconditions": "System is accessible with required test data available",
                    "postconditions": "System is restored to original state; test data cleaned up if necessary",
                    "expected": expected,
                })
                idx += 1
            high_quality_scenarios.append(scenario)

        # Now process brand-specific scenarios
        for scenario_obj in brand_scenarios:
            scenario = scenario_obj["text"]
            specific_brand = scenario_obj["brand"]

            if test_type and test_type.lower() not in scenario.lower():
                continue

            component = None
            if components:
                for comp in components:
                    if comp.lower() in scenario.lower():
                        component = comp
                        break

            # For brand-specific scenarios, use the brand name from the scenario object
            brand_code = self._extract_brands(specific_brand)

            matched_steps = [s for s in steps_list if s in scenario]
            scenario_steps = self._generate_scenario_steps(
                scenario,
                matched_steps if matched_steps else None,
                {"acceptance": acceptance_criteria, "flows": flows}
            )
            automatable = self._determine_automation_feasibility(scenario, scenario_steps)

            for case_type in ["POS", "NEG", "EDGE"]:
                test_data = self._generate_realistic_test_data(scenario, scenario_steps, case_type)
                expected = self._craft_expected_result(scenario, scenario_steps, case_type, acceptance_criteria)
                test_name = self._generate_branded_test_name(scenario, case_type, brand_code)

                # Create brand-specific description
                description = scenario.strip()
                if case_type == "POS":
                    description = f"[{specific_brand.upper()}] This test verifies that {scenario.strip().lower()}"
                elif case_type == "NEG":
                    description = f"[{specific_brand.upper()}] This test verifies proper error handling when {scenario.strip().lower()}"
                else:  # EDGE
                    description = f"[{specific_brand.upper()}] This test validates boundary conditions when {scenario.strip().lower()}"

                tc_id = self._make_test_name(component, scenario, case_type, idx)
                conditions_steps_formatted = '\n'.join(f"{i+1}. {s}" for i, s in enumerate(scenario_steps))

                test_cases.append({
                    "TC id": tc_id,
                    "brand": brand_code,
                    "name": test_name,
                    "description": description,
                    "conditions/steps": conditions_steps_formatted,
                    "test data": test_data,
                    "priority": "High" if case_type == "POS" else ("Medium" if case_type == "EDGE" else "Low"),
                    "severity": "Critical" if case_type == "POS" else ("Major" if case_type == "EDGE" else "Minor"),
                    "automation feasibility": automatable,
                    "expected result": expected,
                    "_automatable": automatable == "Yes",
                    "name_robot": f"{tc_id} {test_name[:40]}",
                    "steps": scenario_steps,
                    "preconditions": f"System is accessible with required test data available for {specific_brand}",
                    "postconditions": "System is restored to original state; test data cleaned up if necessary",
                    "expected": expected,
                })
                idx += 1

            # Add to high-quality scenarios
            high_quality_scenarios.append(scenario)

        # Rest of the function (processing func_items and ambiguous requirements)
        for func in func_items:
            if any(func in scenario for scenario in high_quality_scenarios):
                continue
            if test_type and test_type.lower() not in func.lower():
                continue
            brand = self._extract_brands(func)
            func_steps = self._generate_scenario_steps(
                func,
                [],
                {"acceptance": acceptance_criteria, "flows": flows},
            )
            automatable = self._determine_automation_feasibility(func, func_steps)
            for case_type in ["POS", "NEG", "EDGE"]:
                test_data = self._generate_realistic_test_data(func, func_steps, case_type)
                expected = self._craft_expected_result(func, func_steps, case_type, acceptance_criteria)
                test_name = self._generate_branded_test_name(func, case_type, brand)
                tc_id = self._make_test_name(None, func, case_type, idx)
                conditions_steps = '\n'.join(f"{i+1}. {s}" for i, s in enumerate(func_steps))
                test_cases.append({
                    "TC id": tc_id,
                    "brand": brand,
                    "name": test_name,
                    "description": func,
                    "conditions/steps": conditions_steps,
                    "test data": test_data,
                    "priority": "High" if case_type == "POS" else ("Medium" if case_type == "EDGE" else "Low"),
                    "severity": "Critical" if case_type == "POS" else ("Major" if case_type == "EDGE" else "Minor"),
                    "automation feasibility": automatable,
                    "expected result": expected,
                    "_automatable": automatable == "Yes",
                    "name_robot": f"{tc_id} {test_name[:40]}",
                    "steps": func_steps,
                    "preconditions": "System is accessible with required test data available",
                    "postconditions": "System is restored to original state; test data cleaned up if necessary",
                    "expected": expected,
                })
                idx += 1
        for amb in ambiguous:
            if len(amb.strip()) < 15:
                continue
            tc_id = f"TC_AMBIGUOUS_{str(idx).zfill(3)}"
            amb_steps = [
                "Clarify the ambiguous requirement with business/product team",
                "Document the clarified requirement",
                "Update test cases based on clarification"
            ]
            test_cases.append({
                "TC id": tc_id,
                "brand": "All Brands",
                "name": f"Ambiguity - {amb[:60]}...",
                "description": f"REQUIRES CLARIFICATION: {amb}",
                "conditions/steps": '\n'.join(f"{i+1}. {s}" for i, s in enumerate(amb_steps)),
                "test data": "N/A - Requirement needs clarification before test data can be determined",
                "priority": "Medium",
                "severity": "Major",
                "automation feasibility": "No",
                "expected result": "Requirement is clarified and documented properly for testing",
                "_automatable": False,
                "name_robot": f"{tc_id} {amb[:40]}",
                "steps": amb_steps,
                "preconditions": "",
                "postconditions": "",
                "expected": "Ambiguous requirement is clarified and properly documented",
            })
            idx += 1
        return test_cases


def generate_test_plan(analysis, config=None, scope=None, test_type=None, components=None):
    """
    Generate a comprehensive test plan document in Word format based on the analyzed requirements.

    Args:
        analysis: The requirement analysis results
        config: Configuration parameters
        scope: Test scope (Functional, Non-Functional, or Both)
        test_type: Test type (Component, Integration, Acceptance, or All)
        components: Target components/modules

    Returns:
        Filepath to the generated test plan document
    """
    import time
    start_time = time.time()
    MAX_TEST_CASES = 30  # Maximum number of test cases to include in the document
    TIMEOUT_SECONDS = 120  # Maximum time to spend on test case generation

    # Create a new Word document
    doc = Document()

    # Add title
    doc.add_heading('Test Plan Document', 0)

    # Add metadata section
    doc.add_heading('1. Test Plan Information', level=1)
    doc.add_paragraph(f'Generated Date: {datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph(f'Test Scope: {scope or "Both"}')
    doc.add_paragraph(f'Test Type: {test_type or "All"}')
    doc.add_paragraph(f'Target Components: {", ".join(components) if components else "All Components"}')

    # Add introduction
    doc.add_heading('2. Introduction', level=1)
    doc.add_paragraph(analysis.get('system_description', 'No system description available.'))

    # Add test scope and objectives
    doc.add_heading('3. Test Scope', level=1)
    if scope == "Functional":
        doc.add_paragraph('This test plan covers functional testing of the system, focusing on verifying that the application functions correctly according to the specified requirements.')
    elif scope == "Non-Functional":
        doc.add_paragraph('This test plan covers non-functional testing aspects such as performance, security, usability, reliability, and scalability.')
    else:
        doc.add_paragraph('This test plan covers both functional and non-functional testing aspects of the system.')

    # Add test objectives
    doc.add_heading('4. Test Objectives', level=1)
    doc.add_paragraph('The objectives of this testing effort are:')
    objectives = [
        "Verify that the system meets all functional and business requirements",
        "Identify defects and issues early in the development lifecycle",
        "Ensure the quality and reliability of the system",
        "Validate that the application performs well under expected conditions",
        "Confirm that the system is secure and protected against vulnerabilities"
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style='List Bullet')

    # Add test strategy
    doc.add_heading('5. Test Strategy', level=1)

    if test_type == "Component":
        doc.add_paragraph('Testing will focus on individual components/modules to ensure they function correctly in isolation.')
        doc.add_heading('Component Testing Strategy', level=2)
        doc.add_paragraph('Component testing will be performed to verify that individual units of code function as expected.')
    elif test_type == "Integration":
        doc.add_paragraph('Testing will focus on interactions between components and systems to ensure they work together correctly.')
        doc.add_heading('Integration Testing Strategy', level=2)
        doc.add_paragraph('Integration testing will validate the interfaces and data flow between interconnected components.')
    elif test_type == "Acceptance":
        doc.add_paragraph('Testing will focus on validating that the system meets acceptance criteria and business requirements.')
        doc.add_heading('Acceptance Testing Strategy', level=2)
        doc.add_paragraph('User acceptance testing will be conducted to ensure the system satisfies the business requirements and is ready for production.')
    else:
        doc.add_paragraph('This test plan includes a comprehensive testing strategy covering component, integration, and acceptance testing phases.')
        doc.add_heading('Multi-level Testing Strategy', level=2)
        doc.add_paragraph('Testing will be conducted at multiple levels, from unit/component testing through integration testing to system and acceptance testing.')

    # Add test environment
    doc.add_heading('6. Test Environment', level=1)
    doc.add_paragraph('The following environments will be used for testing:')
    environments = [
        "Development Environment: For component and early integration testing",
        "QA/Testing Environment: For comprehensive testing and test automation",
        "Staging Environment: For final validation before production deployment",
        "Production-like Environment: For performance and load testing"
    ]
    for env in environments:
        doc.add_paragraph(env, style='List Bullet')

    # Add entry and exit criteria
    doc.add_heading('7. Entry and Exit Criteria', level=1)

    doc.add_heading('7.1 Entry Criteria', level=2)
    entry_criteria = [
        "Requirements have been reviewed and approved",
        "Test environment is set up and configured",
        "Test data is prepared and available",
        "Test cases have been developed and reviewed",
        "Development team has completed the initial build",
        "Critical path functionality is implemented"
    ]
    for criterion in entry_criteria:
        doc.add_paragraph(criterion, style='List Bullet')

    doc.add_heading('7.2 Exit Criteria', level=2)
    exit_criteria = [
        "All planned test cases have been executed",
        "All critical defects have been resolved or deferred",
        "Test coverage meets the defined criteria",
        "Test results have been reviewed and approved",
        "Test environment is cleaned up and restored",
        "Test documentation is complete and archived"
    ]

    for criterion in exit_criteria:
        doc.add_paragraph(criterion, style='List Bullet')

    # Add test deliverables
    doc.add_heading('8. Test Deliverables', level=1)
    doc.add_paragraph('The following deliverables will be produced as part of the testing process:')
    deliverables = [
        "Test Plan Document",
        "Test Case Specifications",
        "Test Data Sets",
        "Test Execution Results",
        "Defect Reports",
        "Test Summary Report"
    ]
    for deliverable in deliverables:
        doc.add_paragraph(deliverable, style='List Bullet')

    # Add test schedule
    doc.add_heading('9. Test Schedule', level=1)
    doc.add_paragraph('The testing schedule will be aligned with the overall project timeline. Key milestones include:')
    milestones = [
        "Test Planning: [Start Date] to [End Date]",
        "Test Case Development: [Start Date] to [End Date]",
        "Test Execution: [Start Date] to [End Date]",
        "Test Closure: [Start Date] to [End Date]"
    ]
    for milestone in milestones:
        doc.add_paragraph(milestone, style='List Bullet')

    # Add risk assessment
    doc.add_heading('10. Risk Assessment', level=1)
    doc.add_paragraph('The following risks have been identified that may impact the testing process:')
    risks = [
        "Inadequate test coverage due to incomplete requirements",
        "Delays in development impacting testing timelines",
        "Resource availability and skill gaps",
        "Environment setup issues",
        "Defects found late in the testing cycle"
    ]
    for risk in risks:
        doc.add_paragraph(risk, style='List Bullet')

    # Add test case generation with limits and progress tracking
    doc.add_heading('11. Test Case Generation', level=1)
    doc.add_paragraph('The following representative test cases have been generated based on the analysis:')
    doc.add_paragraph('Note: A maximum of 30 test cases are shown for document readability. Additional test cases can be generated separately.')

    print(f"Generating test cases with a limit of {MAX_TEST_CASES}...")
    test_case_generator = TestCaseGenerator(analysis, config)

    # Initialize test_cases to empty list to avoid reference errors
    test_cases = []

    try:
        # Set a timeout for test case generation
        generation_start = time.time()

        # Generate test cases with timeout
        while len(test_cases) == 0 and (time.time() - generation_start) < TIMEOUT_SECONDS:
            test_cases = test_case_generator.generate(scope, test_type, components)
            if not test_cases:
                print("No test cases generated yet, retrying...")
                time.sleep(1)

        print(f"Generated {len(test_cases)} test cases in {time.time() - generation_start:.2f} seconds")

        # Limit the number of test cases to include in the document
        if len(test_cases) > MAX_TEST_CASES:
            print(f"Limiting to {MAX_TEST_CASES} test cases in the document")
            # Prioritize: Include some POS, NEG, and EDGE cases in a balanced way
            pos_cases = [tc for tc in test_cases if "POS" in tc.get('TC id', '')]
            neg_cases = [tc for tc in test_cases if "NEG" in tc.get('TC id', '')]
            edge_cases = [tc for tc in test_cases if "EDGE" in tc.get('TC id', '')]

            # Calculate how many of each type to include
            total_slots = MAX_TEST_CASES
            pos_slots = min(len(pos_cases), total_slots // 2)  # 50% positive cases
            neg_slots = min(len(neg_cases), total_slots // 3)  # ~33% negative cases
            edge_slots = min(len(edge_cases), total_slots - pos_slots - neg_slots)  # remainder for edge cases

            # Adjust if we have slots left
            remaining_slots = total_slots - pos_slots - neg_slots - edge_slots
            if remaining_slots > 0:
                if len(pos_cases) > pos_slots:
                    pos_slots += remaining_slots
                elif len(neg_cases) > neg_slots:
                    neg_slots += remaining_slots
                elif len(edge_cases) > edge_slots:
                    edge_slots += remaining_slots

            # Select the cases
            selected_cases = (
                pos_cases[:pos_slots] +
                neg_cases[:neg_slots] +
                edge_cases[:edge_slots]
            )
            test_cases = selected_cases[:MAX_TEST_CASES]

        # Add the test cases to the document
        for i, test_case in enumerate(test_cases):
            doc.add_paragraph(f"Test Case {i+1} of {len(test_cases)}")
            doc.add_paragraph(f"Test Case ID: {test_case['TC id']}")
            doc.add_paragraph(f"Name: {test_case['name']}")
            doc.add_paragraph(f"Description: {test_case['description']}")
            doc.add_paragraph(f"Steps: {test_case['conditions/steps']}")
            doc.add_paragraph(f"Priority: {test_case['priority']}")
            doc.add_paragraph(f"Severity: {test_case['severity']}")
            doc.add_paragraph(f"Automation Feasibility: {test_case['automation feasibility']}")
            doc.add_paragraph(f"Expected Result: {test_case['expected result']}")
            doc.add_paragraph("\n")

            # Save progress periodically to avoid memory issues with large documents
            if i > 0 and i % 10 == 0:
                print(f"Added {i} test cases to document, saving progress...")
                temp_output = os.path.join(os.getcwd(), f'test_plan_progress_{i}.docx')
                doc.save(temp_output)

    except Exception as e:
        # In case of error, add information to the document and continue
        print(f"Error during test case generation: {e}")
        doc.add_paragraph(f"Error during test case generation: {str(e)}")
        doc.add_paragraph("Please try again with a smaller scope or different parameters.")

    # Add total generation time
    total_time = time.time() - start_time
    doc.add_paragraph(f"Test plan generation completed in {total_time:.2f} seconds.")

    # Provide more detailed statistics about the generated test cases
    if test_cases:
        doc.add_heading("Test Coverage Statistics", level=2)
        pos_cases = len([tc for tc in test_cases if "POS" in tc.get('TC id', '')])
        neg_cases = len([tc for tc in test_cases if "NEG" in tc.get('TC id', '')])
        edge_cases = len([tc for tc in test_cases if "EDGE" in tc.get('TC id', '')])
        automatable = len([tc for tc in test_cases if tc.get('_automatable', False)])

        coverage_table = doc.add_table(rows=1, cols=2)
        coverage_table.style = 'Table Grid'
        hdr_cells = coverage_table.rows[0].cells
        hdr_cells[0].text = 'Metric'
        hdr_cells[1].text = 'Count'

        metrics = [
            ('Total Test Cases', len(test_cases)),
            ('Positive Test Cases', pos_cases),
            ('Negative Test Cases', neg_cases),
            ('Edge Cases', edge_cases),
            ('Automatable Test Cases', automatable),
            ('Brands Covered', len(set(tc.get('brand', 'Unknown') for tc in test_cases))),
        ]

        for metric, count in metrics:
            row_cells = coverage_table.add_row().cells
            row_cells[0].text = metric
            row_cells[1].text = str(count)

    # Save the final document with timestamp to avoid overwriting
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    output_dir = os.path.join(os.getcwd(), config.get('output_dir', 'generated_tests'))
    os.makedirs(output_dir, exist_ok=True)
    output_file = os.path.join(output_dir, f'test_plan_document_{timestamp}.docx')

    try:
        doc.save(output_file)
        print(f"Test plan document generated: {output_file}")
    except Exception as save_error:
        # Fallback to saving in current directory if output directory is not writable
        fallback_output = os.path.join(os.getcwd(), f'test_plan_document_{timestamp}.docx')
        doc.save(fallback_output)
        print(f"Error saving to specified output directory. Test plan saved to: {fallback_output}")
        output_file = fallback_output

    return output_file

